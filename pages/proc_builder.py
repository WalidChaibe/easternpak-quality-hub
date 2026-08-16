import streamlit as st
import pandas as pd
from datetime import date, timedelta
import base64
import requests
import io
import math
import re
import json
import subprocess
import tempfile
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.units import cm, mm
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    Image, HRFlowable, PageBreak, KeepTogether
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY
from reportlab.platypus.flowables import Flowable
from reportlab.pdfbase.pdfmetrics import stringWidth

from utils.auth import require_auth, can_write, get_profile
from utils.supabase_client import get_supabase

# ─────────────────────────────────────────────
# CONSTANTS
# ─────────────────────────────────────────────
LOGO_PATH  = "static/napco_logo.png"
NAPCO_BLUE = colors.HexColor("#0D68A3")
LIGHT_BLUE = colors.HexColor("#D5E8F0")
DARK_TEXT  = colors.HexColor("#1F2937")
GRAY_BG    = colors.HexColor("#F5F5F5")
PAGE_W, PAGE_H = A4
LEFT_M = RIGHT_M = 2*cm
TOP_M  = 3.2*cm   # space for header
BOT_M  = 2*cm

DEPT_MAP = {
    "SC — Supply Chain":         {"code": "SC",  "subs": ["CS","MH","AW","WH","XX","Other"]},
    "PL — Plant":                {"code": "PL",  "subs": ["MT","PD","XX","Other"]},
    "TQA — Technical & Quality": {"code": "TQA", "subs": ["QC","TQ","XX","Other"]},
    "PM — Product Management":   {"code": "PM",  "subs": ["XX","Other"]},
    "Other":                     {"code": None,  "subs": ["XX","Other"]},
}
DOC_TYPE_MAP = {"PD — Procedure": "PD", "PR — Process": "PR"}
SWIMLANES_DEFAULT = ["Step Owner","QC","Production","Maintenance","Management","HSE","Supply Chain"]

# ─────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────
def ordinal(n):
    s = {1:"st",2:"nd",3:"rd"}
    return f"{n}{s.get(n%10 if n%100 not in (11,12,13) else 0,'th')}"

def revision_label(rev):
    return "1st Issue" if rev == 0 else f"{ordinal(rev)} Revision"

def build_mermaid(steps, lanes):
    if not steps: return ""
    lines = ["flowchart TD"]
    lane_steps = {}
    for s in steps:
        lane = s.get("swimlane","General")
        lane_steps.setdefault(lane, [])
        lane_steps[lane].append(s)
    for lane, ls in lane_steps.items():
        safe = lane.replace(" ","_").replace("-","_").replace("/","_")
        lines.append(f'    subgraph {safe}["{lane}"]')
        for s in ls:
            sid   = s["id"].replace("-","")
            label = s.get("title","Step").replace('"',"'")
            shape = s.get("shape","rect")
            if shape == "diamond":
                lines.append(f'        {sid}{{{{{label}}}}}')
            elif shape == "rounded":
                lines.append(f'        {sid}([{label}])')
            else:
                lines.append(f'        {sid}[{label}]')
        lines.append("    end")
    for i, s in enumerate(steps[:-1]):
        src  = steps[i]["id"].replace("-","")
        dst  = steps[i+1]["id"].replace("-","")
        conn = s.get("connection_label","")
        lines.append(f"    {src} -->|{conn}| {dst}" if conn else f"    {src} --> {dst}")
    return "\n".join(lines)

_CATEGORY_MAP = {
    "PD": "Procedure", "PR": "Procedure", "PROCEDURE": "Procedure",
    "FM": "Form", "FORM": "Form",
    "WI": "Work Instruction", "WORK INSTRUCTION": "Work Instruction",
    "PO": "Policy", "POL": "Policy", "POLICY": "Policy",
}

def _normalize_category(doc_type):
    """Map internal doc_type abbreviations (PD, FM, WI, ...) to the same
    human-readable category used everywhere else, so the same kind of
    document never displays under two different labels."""
    if not doc_type:
        return "Other"
    return _CATEGORY_MAP.get(doc_type.strip().upper(), doc_type)

def check_or_add_master(sb, code, title, doc_type, is_internal, uid):
    if not code and not title: return
    key = code or title
    res = sb.table("master_documents").select("id").eq("doc_code", key).execute()
    if not res.data:
        try:
            sb.table("master_documents").insert({
                "doc_code": key, "title": title or key,
                "doc_type": doc_type, "category": _normalize_category(doc_type),
                "is_internal": is_internal, "created_by": uid,
            }).execute()
        except Exception:
            pass


# ─────────────────────────────────────────────
# REPORTLAB SWIMLANE FLOWCHART FLOWABLE
# ─────────────────────────────────────────────

def _clean_approval_label(text):
    """Some legacy documents embed '(Prepared by)' / '(Reviewed by)' /
    '(Approved by)' directly inside the function name. The approvals table
    already has a 'Prepared by / Reviewed by / Approved by' style role for
    each column via its position, so this suffix is redundant — strip it."""
    if not text:
        return text
    return re.sub(r"\s*\((?:prepared|reviewed|approved)\s+by\)\s*", "",
                   text, flags=re.IGNORECASE).strip()

def _draw_filled_poly(canvas, points, fill_color, stroke_color):
    p = canvas.beginPath()
    p.moveTo(points[0], points[1])
    for i in range(2, len(points), 2):
        p.lineTo(points[i], points[i+1])
    p.close()
    canvas.setFillColor(fill_color)
    canvas.setStrokeColor(stroke_color)
    canvas.drawPath(p, fill=1, stroke=1)

def _draw_arrowhead(canvas, x, y, color):
    p = canvas.beginPath()
    p.moveTo(x, y)
    p.lineTo(x - 0.12*cm, y + 0.22*cm)
    p.lineTo(x + 0.12*cm, y + 0.22*cm)
    p.close()
    canvas.setFillColor(color)
    canvas.setStrokeColor(color)
    canvas.drawPath(p, fill=1, stroke=0)

def _draw_arrowhead_right(canvas, x, y, color):
    p = canvas.beginPath()
    p.moveTo(x, y)
    p.lineTo(x - 0.22*cm, y + 0.12*cm)
    p.lineTo(x - 0.22*cm, y - 0.12*cm)
    p.close()
    canvas.setFillColor(color)
    canvas.setStrokeColor(color)
    canvas.drawPath(p, fill=1, stroke=0)

def _wrap_words(text, font, size, max_width):
    """Greedy word-wrap: returns a list of lines that each fit within max_width.
    A single space-delimited "word" that's itself too wide (e.g. a long
    slash-separated term like Tools/Equipment/Accessories) is additionally
    split on '/' as a fallback break point, so it doesn't overflow its box.
    Each (piece, glue) pair tracks whether the piece joins the previous
    piece with no space (glue='') or a normal space (glue=' ')."""
    raw_words = (text or "").split()
    if not raw_words:
        return [""]

    pieces = []  # list of (text, glue_before)
    for w in raw_words:
        if stringWidth(w, font, size) > max_width and "/" in w:
            sub = w.split("/")
            for i, p in enumerate(sub):
                token = p + ("/" if i < len(sub) - 1 else "")
                pieces.append((token, "" if i > 0 else " "))
        else:
            pieces.append((w, " "))

    lines, cur_text, cur_pieces = [], "", []
    for token, glue in pieces:
        candidate = cur_text + glue + token if cur_pieces else token
        if stringWidth(candidate, font, size) <= max_width or not cur_pieces:
            cur_text = candidate
            cur_pieces.append(token)
        else:
            lines.append(cur_text)
            cur_text = token
            cur_pieces = [token]
    if cur_pieces:
        lines.append(cur_text)
    return lines

class SwimlaneFlowchart(Flowable):
    """Draws a top-down process flowchart with support for:
      - a horizontal row of parallel "trigger" boxes that fan into one point
        (step type "input_row"), matching multiple inputs feeding one step
      - a horizontal side-branch off any step (e.g. a parallel procedure
        that gets triggered), drawn to the right with its own arrow
      - a decision that splits into two (or more) parallel multi-step
        branches that later merge back into a single point (step type
        "branch_split"), e.g. "Outsource Calibration? Yes/No" each leading
        down their own short sequence before rejoining the main flow

    Everything else is a single centered step column, same as before.
    Box/row heights are computed up front from the actual wrapped text, so
    nothing overflows regardless of title length.
    """

    BOX_W        = 7.6*cm
    MIN_BOX_H    = 1.0*cm
    DIA_W        = 6.2*cm
    MIN_DIA_H    = 1.7*cm
    V_GAP        = 1.05*cm     # vertical gap between steps (room for arrow + label)
    ROLE_LBL_H   = 0.55*cm     # space reserved above each box for the role chip
    LANE_PAD     = 0.5*cm      # top/bottom padding of the whole chart
    LANE_HDR_H   = 0           # no header row in the single-column layout
    FONT         = "Helvetica"
    FONT_B       = "Helvetica-Bold"
    FONT_SZ      = 8
    ROLE_FONT_SZ = 6.3
    LINE_H       = FONT_SZ + 2.4

    BRANCH_COL_W  = 5.0*cm     # width of each column inside a branch_split
    BRANCH_GAP    = 0.5*cm     # gap between branch columns
    BRANCH_V_GAP  = 0.55*cm    # vertical gap between boxes within a branch column
    BRANCH_FONT_SZ= 7.3
    BRANCH_LBL_H  = 0.4*cm     # space for the branch label ("YES - External") above its column

    BOX_FILL     = colors.HexColor("#EAF3FB")
    BOX_BORDER   = NAPCO_BLUE
    DIA_FILL     = colors.HexColor("#FFF6D8")
    DIA_BORDER   = colors.HexColor("#C9A227")
    BRANCH_FILL  = colors.HexColor("#FCEAEA")
    BRANCH_BORDER= colors.HexColor("#B03A3A")
    SHADOW_COLOR = colors.HexColor("#DADADA")
    ARROW_COLOR  = colors.HexColor("#4A4A4A")

    def __init__(self, steps, available_width, all_lanes=None):
        # all_lanes accepted for backward compatibility; unused.
        super().__init__()
        self.steps   = steps
        self.avail_w = available_width
        self.width   = available_width

        # If any step has a side-branch, shift the main column left so
        # there's room on the right for the branch box.
        self._has_branch = any(s.get("side_branch") for s in steps
                                if s.get("type", "step") == "step")
        self.cx = self.width * 0.36 if self._has_branch else self.width / 2

        self._prepared = []
        for step in steps:
            stype = step.get("type", "step")

            if stype == "input_row":
                items = step.get("items", [])
                n = max(1, len(items))
                gap = 0.22*cm
                item_w = min(3.4*cm, (self.width - gap*(n-1) - 2*self.LANE_PAD) / n)
                prepared_items, max_lines = [], 1
                for it in items:
                    lines = _wrap_words(it.get("title",""), self.FONT,
                                         self.ROLE_FONT_SZ + 0.7, item_w - 0.3*cm)
                    prepared_items.append(lines)
                    max_lines = max(max_lines, len(lines))
                row_h = max(1.0*cm, max_lines * (self.ROLE_FONT_SZ + 2.6) + 0.25*cm)
                bus_h = 0.6*cm  # room for the converging arrows below the row
                self._prepared.append({
                    "type": "input_row", "items": prepared_items, "item_w": item_w,
                    "row_h": row_h, "box_h": row_h + bus_h, "bus_h": bus_h,
                })
                continue

            if stype == "branch_split":
                branches = step.get("branches", [])
                prepared_branches = []
                for br in branches:
                    br_steps = []
                    col_h = self.BRANCH_LBL_H
                    for s in br.get("steps", []):
                        lines = _wrap_words(s.get("title",""), self.FONT,
                                             self.BRANCH_FONT_SZ, self.BRANCH_COL_W - 0.6*cm)
                        text_h = len(lines) * (self.BRANCH_FONT_SZ + 2.2)
                        box_h  = max(0.9*cm, text_h + 0.3*cm)
                        br_steps.append({"lines": lines, "box_h": box_h})
                        col_h += box_h + self.BRANCH_V_GAP
                    prepared_branches.append({
                        "label": br.get("label",""), "steps": br_steps, "col_h": col_h,
                    })
                max_col_h = max((b["col_h"] for b in prepared_branches), default=0)
                bus_h = 0.6*cm  # room for the converging arrows below the columns
                self._prepared.append({
                    "type": "branch_split", "branches": prepared_branches,
                    "max_col_h": max_col_h, "bus_h": bus_h, "box_h": max_col_h + bus_h,
                })
                continue

            shape = step.get("shape", "rect")
            title = step.get("title", "")
            max_w = (self.DIA_W - 1.3*cm) if shape == "diamond" else (self.BOX_W - 0.7*cm)
            lines  = _wrap_words(title, self.FONT, self.FONT_SZ, max_w)
            text_h = len(lines) * self.LINE_H
            if shape == "diamond":
                box_h = max(self.MIN_DIA_H, text_h + 0.7*cm)
            else:
                box_h = max(self.MIN_BOX_H, text_h + 0.35*cm)

            side_prepared = None
            side = step.get("side_branch")
            if side:
                side_w = self.width - (self.cx + self.BOX_W/2) - 1.1*cm
                side_w = max(2.6*cm, min(side_w, 5.6*cm))
                side_lines = _wrap_words(side.get("title",""), self.FONT,
                                          self.FONT_SZ - 0.5, side_w - 0.5*cm)
                side_h = max(0.95*cm, len(side_lines) * self.LINE_H + 0.3*cm)
                side_prepared = {"lines": side_lines, "w": side_w, "box_h": side_h}

            self._prepared.append({
                "type": "step", "shape": shape, "lines": lines, "box_h": box_h,
                "side_branch": side_prepared,
            })

        self._step_heights = []
        for p in self._prepared:
            lbl_h = self.ROLE_LBL_H if p["type"] == "step" else 0
            self._step_heights.append(lbl_h + p["box_h"] + self.V_GAP)

        self.total_h = self.LANE_PAD + sum(self._step_heights) + self.LANE_PAD
        self.height  = self.total_h

    def _step_geom(self, idx):
        """Returns (center_y, box_h) for step idx, from the bottom of the flowable."""
        y_from_top = self.LANE_PAD + sum(self._step_heights[:idx])
        p = self._prepared[idx]
        lbl_h = self.ROLE_LBL_H if p["type"] == "step" else 0
        y_from_top += lbl_h + p["box_h"] / 2
        return self.total_h - y_from_top, p["box_h"]

    def _draw_box(self, c, cx, cy, bw, bh, fill, border, radius=8):
        c.setFillColor(self.SHADOW_COLOR)
        c.roundRect(cx - bw/2 + 0.08*cm, cy - bh/2 - 0.08*cm, bw, bh, radius=radius, fill=1, stroke=0)
        c.setFillColor(fill)
        c.setStrokeColor(border)
        c.setLineWidth(1.2)
        c.roundRect(cx - bw/2, cy - bh/2, bw, bh, radius=radius, fill=1, stroke=1)

    def _draw_role_chip(self, c, cx, top_y, role):
        role_txt = role.upper()
        c.setFont(self.FONT_B, self.ROLE_FONT_SZ)
        chip_w = stringWidth(role_txt, self.FONT_B, self.ROLE_FONT_SZ) + 0.34*cm
        chip_h = 0.36*cm
        chip_y = top_y + 0.09*cm
        c.setFillColor(NAPCO_BLUE)
        c.roundRect(cx - chip_w/2, chip_y, chip_w, chip_h, radius=chip_h/2, fill=1, stroke=0)
        c.setFillColor(colors.white)
        c.drawCentredString(cx, chip_y + 0.10*cm, role_txt)

    def _draw_centered_lines(self, c, cx, cy, lines, font, size, color=None):
        c.setFillColor(color or DARK_TEXT)
        c.setFont(font, size)
        line_h = size + 2.4
        start_y = cy + (len(lines) - 1) * line_h / 2
        for li, ln in enumerate(lines):
            c.drawCentredString(cx, start_y - li * line_h, ln)

    def draw(self):
        c  = self.canv
        cx = self.cx
        c.setLineCap(1)
        c.setLineJoin(1)

        c.setStrokeColor(colors.HexColor("#DDDDDD"))
        c.setLineWidth(0.6)
        c.roundRect(1, 1, self.width - 2, self.total_h - 2, radius=6, fill=0, stroke=1)

        for idx, step in enumerate(self.steps):
            p  = self._prepared[idx]
            cy, bh = self._step_geom(idx)
            next_is_branch_split = (idx < len(self.steps) - 1
                                     and self._prepared[idx + 1]["type"] == "branch_split")

            # ── Fan-in input row ──────────────────────────────
            if p["type"] == "input_row":
                item_w, row_h, bus_h = p["item_w"], p["row_h"], p["bus_h"]
                n = len(p["items"])
                gap = 0.22*cm
                total_w = n * item_w + (n - 1) * gap
                start_x = (self.width - total_w) / 2
                row_cy  = cy + bus_h / 2

                centers = []
                for i, lines in enumerate(p["items"]):
                    bx = start_x + i * (item_w + gap)
                    bx_c = bx + item_w / 2
                    centers.append(bx_c)
                    self._draw_box(c, bx_c, row_cy, item_w, row_h,
                                    self.BOX_FILL, self.BOX_BORDER, radius=6)
                    self._draw_centered_lines(c, bx_c, row_cy, lines,
                                               self.FONT, self.ROLE_FONT_SZ + 0.7)

                bus_y = row_cy - row_h / 2 - bus_h * 0.5
                c.setStrokeColor(self.ARROW_COLOR)
                c.setLineWidth(1.0)
                for bx_c in centers:
                    c.line(bx_c, row_cy - row_h / 2, bx_c, bus_y)
                c.line(min(centers), bus_y, max(centers), bus_y)

                if idx < len(self.steps) - 1:
                    next_cy, next_bh = self._step_geom(idx + 1)
                    next_top = next_cy + next_bh / 2
                    y_end = next_top + 0.06*cm
                    c.setFillColor(self.ARROW_COLOR)
                    c.line(cx, bus_y, cx, y_end)
                    _draw_arrowhead(c, cx, y_end, self.ARROW_COLOR)
                continue

            # ── Branch split: decision fans out into parallel columns, ──
            # ── each column runs its own steps, then all columns merge ──
            if p["type"] == "branch_split":
                branches  = p["branches"]
                n         = len(branches)
                gap       = self.BRANCH_GAP
                total_w   = n * self.BRANCH_COL_W + (n - 1) * gap
                start_x   = cx - total_w / 2
                block_top = cy + bh / 2   # top of this whole block (right below the decision)

                # Fan-out from the previous step's bottom into a shared bus,
                # then down into each column.
                col_centers = [start_x + i * (self.BRANCH_COL_W + gap) + self.BRANCH_COL_W / 2
                               for i in range(n)]
                if idx > 0:
                    prev_cy, prev_bh = self._step_geom(idx - 1)
                    prev_bottom = prev_cy - prev_bh / 2
                    fanout_bus_y = block_top - 0.15*cm
                    c.setStrokeColor(self.ARROW_COLOR)
                    c.setLineWidth(1.0)
                    c.line(cx, prev_bottom, cx, fanout_bus_y)
                    c.line(min(col_centers), fanout_bus_y, max(col_centers), fanout_bus_y)
                    for col_cx in col_centers:
                        y_end = block_top - self.BRANCH_LBL_H + 0.06*cm
                        c.line(col_cx, fanout_bus_y, col_cx, y_end)
                        _draw_arrowhead(c, col_cx, y_end, self.ARROW_COLOR)

                col_bottoms = []
                for bi, br in enumerate(branches):
                    col_cx = col_centers[bi]
                    c.setFont(self.FONT_B, self.ROLE_FONT_SZ)
                    c.setFillColor(self.ARROW_COLOR)
                    c.drawCentredString(col_cx, block_top - self.BRANCH_LBL_H + 0.12*cm, br["label"])

                    y_cursor = block_top - self.BRANCH_LBL_H
                    prev_box_bottom = None
                    for s in br["steps"]:
                        box_h  = s["box_h"]
                        box_cy = y_cursor - box_h / 2
                        self._draw_box(c, col_cx, box_cy, self.BRANCH_COL_W, box_h,
                                        self.BOX_FILL, self.BOX_BORDER, radius=6)
                        self._draw_centered_lines(c, col_cx, box_cy, s["lines"],
                                                   self.FONT, self.BRANCH_FONT_SZ)
                        if prev_box_bottom is not None:
                            c.setStrokeColor(self.ARROW_COLOR)
                            c.setFillColor(self.ARROW_COLOR)
                            c.setLineWidth(1.0)
                            y_end = box_cy + box_h / 2 + 0.05*cm
                            c.line(col_cx, prev_box_bottom, col_cx, y_end)
                            _draw_arrowhead(c, col_cx, y_end, self.ARROW_COLOR)
                        y_cursor -= box_h + self.BRANCH_V_GAP
                        prev_box_bottom = box_cy - box_h / 2
                    col_bottoms.append(prev_box_bottom)

                # Fan-in: converge all column bottoms into the next step.
                bus_y = min(col_bottoms) - (p["bus_h"] * 0.5 if col_bottoms else 0)
                c.setStrokeColor(self.ARROW_COLOR)
                c.setLineWidth(1.0)
                for cb, col_cx in zip(col_bottoms, col_centers):
                    c.line(col_cx, cb, col_cx, bus_y)
                c.line(min(col_centers), bus_y, max(col_centers), bus_y)

                if idx < len(self.steps) - 1:
                    next_cy, next_bh = self._step_geom(idx + 1)
                    next_top = next_cy + next_bh / 2
                    y_end = next_top + 0.06*cm
                    c.setFillColor(self.ARROW_COLOR)
                    c.line(cx, bus_y, cx, y_end)
                    _draw_arrowhead(c, cx, y_end, self.ARROW_COLOR)
                continue

            # ── Regular step (rect / diamond) ─────────────────
            shape = p["shape"]
            lines = p["lines"]
            role  = step.get("swimlane", "")
            conn  = step.get("connection_label", "")

            if shape == "diamond":
                hw, hh = self.DIA_W / 2, bh / 2
                off = 0.08*cm
                shadow_pts = [cx+off, cy+hh-off, cx+hw+off, cy-off,
                              cx+off, cy-hh-off, cx-hw+off, cy-off]
                _draw_filled_poly(c, shadow_pts, self.SHADOW_COLOR, self.SHADOW_COLOR)
                pts = [cx, cy+hh, cx+hw, cy, cx, cy-hh, cx-hw, cy]
                c.setLineWidth(1.2)
                _draw_filled_poly(c, pts, self.DIA_FILL, self.DIA_BORDER)
                box_top, box_bottom = cy + hh, cy - hh
            else:
                self._draw_box(c, cx, cy, self.BOX_W, bh, self.BOX_FILL, self.BOX_BORDER)
                box_top, box_bottom = cy + bh/2, cy - bh/2

            if role:
                self._draw_role_chip(c, cx, box_top, role)

            self._draw_centered_lines(c, cx, cy, lines, self.FONT, self.FONT_SZ)

            # ── Side branch (e.g. a parallel procedure this step triggers) ──
            if p.get("side_branch"):
                sb = p["side_branch"]
                sb_w, sb_h = sb["w"], sb["box_h"]
                sb_x  = cx + self.BOX_W/2 + 1.0*cm
                sb_cx = sb_x + sb_w/2

                self._draw_box(c, sb_cx, cy, sb_w, sb_h,
                                self.BRANCH_FILL, self.BRANCH_BORDER, radius=6)
                self._draw_centered_lines(c, sb_cx, cy, sb["lines"], self.FONT, self.FONT_SZ - 0.5)

                c.setStrokeColor(self.ARROW_COLOR)
                c.setFillColor(self.ARROW_COLOR)
                c.setLineWidth(1.1)
                x_start = cx + self.BOX_W/2
                x_end   = sb_x - 0.05*cm
                c.line(x_start, cy, x_end, cy)
                _draw_arrowhead_right(c, x_end, cy, self.ARROW_COLOR)

            # ── Arrow down to next step (skipped if the next step is a ──
            # ── branch_split — it draws its own fan-out from our bottom) ──
            if idx < len(self.steps) - 1 and not next_is_branch_split:
                next_cy, next_bh = self._step_geom(idx + 1)
                next_top = next_cy + next_bh / 2

                c.setStrokeColor(self.ARROW_COLOR)
                c.setFillColor(self.ARROW_COLOR)
                c.setLineWidth(1.1)
                y_end = next_top + 0.06*cm
                c.line(cx, box_bottom, cx, y_end)
                _draw_arrowhead(c, cx, y_end, self.ARROW_COLOR)

                if conn:
                    mid_y = (box_bottom + y_end) / 2
                    c.setFont(self.FONT_B, self.ROLE_FONT_SZ)
                    label_w = stringWidth(conn, self.FONT_B, self.ROLE_FONT_SZ) + 0.3*cm
                    label_h = 0.34*cm
                    lx = cx + 0.16*cm
                    c.setFillColor(colors.white)
                    c.setStrokeColor(self.ARROW_COLOR)
                    c.setLineWidth(0.6)
                    c.roundRect(lx, mid_y - label_h/2, label_w, label_h,
                                radius=label_h/2, fill=1, stroke=1)
                    c.setFillColor(self.ARROW_COLOR)
                    c.drawCentredString(lx + label_w/2, mid_y - 0.09*cm, conn)


def _should_use_lane_grid(steps):
    """Cross-functional swimlanes read far better than a single vertical
    chain when a document has a small, consistently-reused set of roles
    and no branch/merge points to represent. Docs with branch_split or
    input_row keep the single-column renderer since lane-grid can't
    express branching/merging cleanly."""
    if not steps:
        return False
    if any(s.get("type") in ("branch_split", "input_row") for s in steps):
        return False
    lanes = []
    for s in steps:
        lane = (s.get("swimlane") or "").strip()
        if lane and lane not in lanes:
            lanes.append(lane)
    return 2 <= len(lanes) <= 6


class LaneGridFlowchart(Flowable):
    """Cross-functional swimlane chart: one column per role, steps flow
    top-to-bottom in strict sequence, with connectors routed orthogonally
    between columns whenever the owning role changes."""

    ROW_V_GAP    = 0.5*cm
    ROW_MIN_H    = 1.0*cm
    BOX_PAD_X    = 0.22*cm
    LANE_HDR_H   = 0.9*cm
    LANE_GAP     = 0.12*cm
    TOP_PAD      = 0.3*cm
    BOTTOM_PAD   = 0.3*cm
    FONT         = "Helvetica"
    FONT_B       = "Helvetica-Bold"
    FONT_SZ      = 7.6
    LANE_FONT_SZ = 7.8
    LINE_H       = FONT_SZ + 2.2

    LANE_BANDS   = [colors.HexColor("#EAF3FB"), colors.HexColor("#F4F4F4")]
    BOX_FILL     = colors.white
    BOX_BORDER   = NAPCO_BLUE
    SHADOW_COLOR = colors.HexColor("#DADADA")
    ARROW_COLOR  = colors.HexColor("#4A4A4A")

    def __init__(self, steps, available_width, lane_order=None):
        super().__init__()
        self.steps = steps
        self.width = available_width

        if lane_order:
            self.lanes = list(lane_order)
        else:
            self.lanes = []
            for s in steps:
                lane = (s.get("swimlane") or "General").strip() or "General"
                if lane not in self.lanes:
                    self.lanes.append(lane)

        n = max(1, len(self.lanes))
        self.lane_w = (self.width - (n - 1) * self.LANE_GAP) / n
        self.lane_x = [i * (self.lane_w + self.LANE_GAP) for i in range(n)]
        self.lane_index = {lane: i for i, lane in enumerate(self.lanes)}

        box_w = self.lane_w - 2 * self.BOX_PAD_X
        self._prepared = []
        for s in steps:
            title = s.get("title", "")
            lines = _wrap_words(title, self.FONT, self.FONT_SZ, box_w - 10)
            text_h = len(lines) * self.LINE_H
            box_h = max(self.ROW_MIN_H, text_h + 0.3*cm)
            lane = (s.get("swimlane") or "General").strip() or "General"
            self._prepared.append({"lines": lines, "box_h": box_h, "lane": lane})

        self._row_heights = [p["box_h"] + self.ROW_V_GAP for p in self._prepared]
        self.total_h = self.TOP_PAD + self.LANE_HDR_H + sum(self._row_heights) + self.BOTTOM_PAD
        self.height = self.total_h

    def _row_geom(self, idx):
        y_from_top = self.TOP_PAD + self.LANE_HDR_H + sum(self._row_heights[:idx])
        box_h = self._prepared[idx]["box_h"]
        y_from_top += box_h / 2
        return self.total_h - y_from_top, box_h

    def draw(self):
        c = self.canv
        c.setLineCap(1)
        c.setLineJoin(1)

        body_h = self.total_h - self.LANE_HDR_H - self.TOP_PAD
        for i, lane in enumerate(self.lanes):
            lx = self.lane_x[i]
            c.setFillColor(self.LANE_BANDS[i % len(self.LANE_BANDS)])
            c.rect(lx, self.BOTTOM_PAD, self.lane_w, body_h, fill=1, stroke=0)

            c.setFillColor(NAPCO_BLUE)
            c.rect(lx, self.total_h - self.LANE_HDR_H, self.lane_w, self.LANE_HDR_H, fill=1, stroke=0)
            c.setFillColor(colors.white)
            c.setFont(self.FONT_B, self.LANE_FONT_SZ)
            lane_lines = _wrap_words(lane.upper(), self.FONT_B, self.LANE_FONT_SZ, self.lane_w - 8)
            lh = self.LANE_FONT_SZ + 2
            ly = self.total_h - self.LANE_HDR_H/2 + (len(lane_lines) - 1) * lh / 2
            for li, ln in enumerate(lane_lines):
                c.drawCentredString(lx + self.lane_w/2, ly - li*lh, ln)

        c.setStrokeColor(colors.HexColor("#DDDDDD"))
        c.setLineWidth(0.6)
        c.roundRect(1, 1, self.width - 2, self.total_h - 2, radius=4, fill=0, stroke=1)

        for idx, step in enumerate(self.steps):
            p = self._prepared[idx]
            lane_i = self.lane_index[p["lane"]]
            cx = self.lane_x[lane_i] + self.lane_w / 2
            cy, bh = self._row_geom(idx)
            bw = self.lane_w - 2 * self.BOX_PAD_X

            c.setFillColor(self.SHADOW_COLOR)
            c.roundRect(cx - bw/2 + 0.05*cm, cy - bh/2 - 0.05*cm, bw, bh, radius=5, fill=1, stroke=0)
            c.setFillColor(self.BOX_FILL)
            c.setStrokeColor(self.BOX_BORDER)
            c.setLineWidth(1.1)
            c.roundRect(cx - bw/2, cy - bh/2, bw, bh, radius=5, fill=1, stroke=1)

            c.setFillColor(DARK_TEXT)
            c.setFont(self.FONT, self.FONT_SZ)
            lines = p["lines"]
            start_y = cy + (len(lines) - 1) * self.LINE_H / 2
            for li, ln in enumerate(lines):
                c.drawCentredString(cx, start_y - li * self.LINE_H, ln)

            if idx < len(self.steps) - 1:
                next_p = self._prepared[idx + 1]
                next_lane_i = self.lane_index[next_p["lane"]]
                next_cx = self.lane_x[next_lane_i] + self.lane_w / 2
                next_cy, next_bh = self._row_geom(idx + 1)
                box_bottom = cy - bh / 2
                next_top = next_cy + next_bh / 2

                c.setStrokeColor(self.ARROW_COLOR)
                c.setFillColor(self.ARROW_COLOR)
                c.setLineWidth(1.1)

                if abs(next_cx - cx) < 1:
                    y_end = next_top + 0.06*cm
                    c.line(cx, box_bottom, cx, y_end)
                    _draw_arrowhead(c, cx, y_end, self.ARROW_COLOR)
                else:
                    mid_y = (box_bottom + next_top) / 2
                    c.line(cx, box_bottom, cx, mid_y)
                    c.line(cx, mid_y, next_cx, mid_y)
                    y_end = next_top + 0.06*cm
                    c.line(next_cx, mid_y, next_cx, y_end)
                    _draw_arrowhead(c, next_cx, y_end, self.ARROW_COLOR)

# ─────────────────────────────────────────────
# PDF GENERATION
# ─────────────────────────────────────────────
def generate_pdf(doc):
    buffer    = io.BytesIO()
    avail_w   = PAGE_W - LEFT_M - RIGHT_M

    # Styles
    normal   = ParagraphStyle("N",  fontName="Helvetica",      fontSize=9,  leading=13)
    bold_s   = ParagraphStyle("B",  fontName="Helvetica-Bold", fontSize=9,  leading=13)
    h1s      = ParagraphStyle("H1", fontName="Helvetica-Bold", fontSize=12, leading=16,
                               spaceBefore=10, spaceAfter=4, textColor=NAPCO_BLUE)
    h2s      = ParagraphStyle("H2", fontName="Helvetica-Bold", fontSize=10, leading=13,
                               spaceBefore=6, spaceAfter=3)
    centered = ParagraphStyle("C",  fontName="Helvetica",      fontSize=9,  leading=13,
                               alignment=TA_CENTER)
    centered_b=ParagraphStyle("CB", fontName="Helvetica-Bold", fontSize=9,  leading=13,
                               alignment=TA_CENTER)
    title_s  = ParagraphStyle("T",  fontName="Helvetica-Bold", fontSize=20, leading=26,
                               alignment=TA_CENTER, spaceAfter=8)
    dept_s   = ParagraphStyle("D",  fontName="Helvetica-Bold", fontSize=13, leading=18,
                               alignment=TA_CENTER, spaceAfter=4)
    numbered = ParagraphStyle("NM", fontName="Helvetica",      fontSize=9,  leading=13,
                               leftIndent=24, spaceAfter=3)
    small_i  = ParagraphStyle("SI", fontName="Helvetica-Oblique", fontSize=7, leading=9,
                               alignment=TA_CENTER, textColor=colors.grey)

    doc_code   = doc.get("doc_code","—")
    title_txt  = doc.get("title","")
    dept_label = doc.get("dept_label","") or doc.get("dept","")
    adoption   = str(doc.get("date_of_adoption","—"))
    approvals  = doc.get("approvals") or []
    revisions  = doc.get("_revisions") or []
    issue_date = revisions[0]["revised_date"] if revisions else adoption
    rev_date   = revisions[-1]["revised_date"] if revisions else adoption

    # ── Header/Footer callback ──────────────────────────────
    def header_footer(canvas, doc_obj):
        canvas.saveState()
        w, h = A4
        hdr_x   = LEFT_M
        hdr_w   = w - LEFT_M - RIGHT_M
        # Header table moved down from the very top of the page so the logo
        # (drawn above it) has room and doesn't get clipped by the page edge.
        hdr_top = h - 1.7*cm
        row1_h  = 0.65*cm
        row2_h  = 0.55*cm
        hdr_h   = row1_h + row2_h

        # Logo — sits above header table, top-left
        try:
            canvas.drawImage(LOGO_PATH,
                             hdr_x, hdr_top + 0.15*cm,
                             width=2.6*cm, height=1.0*cm,
                             preserveAspectRatio=True, mask='auto')
        except Exception:
            pass

        # Header table background rows
        # Row 1 background
        canvas.setFillColor(GRAY_BG)
        canvas.rect(hdr_x, hdr_top - row1_h, hdr_w, row1_h, fill=1, stroke=0)
        # Row 2 background
        canvas.setFillColor(colors.white)
        canvas.rect(hdr_x, hdr_top - row1_h - row2_h, hdr_w, row2_h, fill=1, stroke=0)

        # Outer border
        canvas.setStrokeColor(colors.HexColor("#999999"))
        canvas.setLineWidth(0.5)
        canvas.rect(hdr_x, hdr_top - hdr_h, hdr_w, hdr_h, fill=0, stroke=1)

        # Column dividers — row 1: TITLE | (title spans) | NBR OF PAGES | page
        col1_w = 2.2*cm
        col3_w = 2.8*cm
        col4_w = 2.2*cm
        col2_w = hdr_w - col1_w - col3_w - col4_w

        # Vertical lines row 1
        canvas.line(hdr_x + col1_w,
                    hdr_top - row1_h, hdr_x + col1_w, hdr_top)
        canvas.line(hdr_x + col1_w + col2_w,
                    hdr_top - row1_h, hdr_x + col1_w + col2_w, hdr_top)
        canvas.line(hdr_x + col1_w + col2_w + col3_w,
                    hdr_top - row1_h, hdr_x + col1_w + col2_w + col3_w, hdr_top)

        # Horizontal divider between rows
        canvas.line(hdr_x, hdr_top - row1_h,
                    hdr_x + hdr_w, hdr_top - row1_h)

        # Row 1 text
        y1 = hdr_top - row1_h + 0.18*cm
        canvas.setFont("Helvetica-Bold", 7.5)
        canvas.setFillColor(colors.black)
        canvas.drawCentredString(hdr_x + col1_w/2, y1, "TITLE")
        canvas.setFont("Helvetica-Bold", 8.5)
        canvas.drawCentredString(hdr_x + col1_w + col2_w/2, y1, title_txt.upper())
        canvas.setFont("Helvetica-Bold", 7.5)
        canvas.drawCentredString(hdr_x + col1_w + col2_w + col3_w/2, y1, "NBR. OF PAGES")
        canvas.setFont("Helvetica", 7.5)
        canvas.drawCentredString(hdr_x + col1_w + col2_w + col3_w + col4_w/2,
                                 y1, f"PAGE {doc_obj.page}")

        # Row 2 columns: CONTROL NBR | doc_code | 1ST ISSUE DATE | date | REVISION DATE | date
        r2_cols = [2.2*cm, 3.5*cm, 2.5*cm, 2.2*cm, 2.5*cm]
        r2_last = hdr_w - sum(r2_cols)
        r2_cols.append(r2_last)
        r2_labels = ["CONTROL NBR.", doc_code,
                     "1ST ISSUE DATE", str(issue_date),
                     "REVISION DATE", str(rev_date)]
        r2_bold   = [True, False, True, False, True, False]
        x_cur = hdr_x
        y2 = hdr_top - hdr_h + 0.15*cm
        for i, (cw, lbl, is_b) in enumerate(zip(r2_cols, r2_labels, r2_bold)):
            if i > 0:
                canvas.setStrokeColor(colors.HexColor("#999999"))
                canvas.setLineWidth(0.5)
                canvas.line(x_cur, hdr_top - row1_h - row2_h,
                            x_cur, hdr_top - row1_h)
            canvas.setFont("Helvetica-Bold" if is_b else "Helvetica", 7)
            canvas.setFillColor(colors.black)
            canvas.drawCentredString(x_cur + cw/2, y2, lbl)
            x_cur += cw

        # Footer — disclaimer text is word-wrapped so it always stays inside the box
        disclaimer = ("THE INFORMATION CONTAINED HEREIN IS PROPRIETARY TO NAPCO NATIONAL AND IT SHALL "
                      "NOT BE USED, REPRODUCED OR DISCLOSED TO OTHERS EXCEPT AS SPECIFICALLY PERMITTED "
                      "IN WRITING BY THE PROPRIETOR.")
        disc_font, disc_sz = "Helvetica-Oblique", 6.5
        disc_max_w  = hdr_w - 0.6*cm
        disc_lines  = _wrap_words(disclaimer, disc_font, disc_sz, disc_max_w)
        line_gap    = 0.30*cm
        footer_h    = 0.25*cm + (len(disc_lines) + 1) * line_gap  # +1 for the UNCONTROLLED line

        fy = 0.5*cm
        canvas.setFillColor(GRAY_BG)
        canvas.rect(LEFT_M, fy, hdr_w, footer_h, fill=1, stroke=1)

        ty = fy + footer_h - 0.28*cm
        canvas.setFont(disc_font, disc_sz)
        canvas.setFillColor(colors.HexColor("#555555"))
        for ln in disc_lines:
            canvas.drawCentredString(w/2, ty, ln)
            ty -= line_gap

        canvas.setFont("Helvetica-BoldOblique", 7)
        canvas.setFillColor(colors.black)
        canvas.drawCentredString(w/2, ty, '"UNCONTROLLED IF PRINTED"')
        canvas.restoreState()

    # ── Build document ──────────────────────────────────────
    pdf = SimpleDocTemplate(
        buffer, pagesize=A4,
        leftMargin=LEFT_M, rightMargin=RIGHT_M,
        topMargin=TOP_M, bottomMargin=BOT_M,
    )
    pdf.onFirstPage  = header_footer
    pdf.onLaterPages = header_footer

    story = []

    # ── COVER PAGE ──────────────────────────────────────────
    story.append(Spacer(1, 0.8*cm))

    # Large logo on cover
    try:
        story.append(Image(LOGO_PATH, width=5*cm, height=2*cm))
    except Exception:
        pass

    story.append(Spacer(1, 0.4*cm))
    story.append(Paragraph(dept_label.upper(), dept_s))
    story.append(Paragraph(title_txt.upper(), title_s))
    story.append(Spacer(1, 0.5*cm))

    # Approvals table
    if approvals:
        n = len(approvals)
        cw = avail_w / (n + 1)
        ap_data = [
            [Paragraph("APPROVED BY", bold_s)] + [""] * n,
            [Paragraph("Department", bold_s)] +
            [Paragraph(a.get("department",""), centered) for a in approvals],
            [Paragraph("Function", bold_s)] +
            [Paragraph(_clean_approval_label(a.get("function","")), centered) for a in approvals],
            [Paragraph("Signature", bold_s)] + [""] * n,
            [Paragraph("Date", bold_s)] + [""] * n,
        ]
        ap_t = Table(ap_data, colWidths=[cw]*(n+1),
                     rowHeights=[0.6*cm, 0.7*cm, 0.7*cm, 1.4*cm, 0.7*cm])
        ap_t.setStyle(TableStyle([
            ("GRID",        (0,0), (-1,-1), 0.5, colors.HexColor("#999999")),
            ("BACKGROUND",  (0,0), (-1,0),  NAPCO_BLUE),
            ("TEXTCOLOR",   (0,0), (-1,0),  colors.white),
            ("BACKGROUND",  (0,1), (0,-1),  LIGHT_BLUE),
            ("FONTNAME",    (0,0), (0,-1),  "Helvetica-Bold"),
            ("FONTSIZE",    (0,0), (-1,-1), 8),
            ("VALIGN",      (0,0), (-1,-1), "MIDDLE"),
            ("ALIGN",       (1,0), (-1,-1), "CENTER"),
            ("SPAN",        (0,0), (-1,0)),
        ]))
        story.append(ap_t)
        story.append(Spacer(1, 0.4*cm))

    # Date of adoption + last reviewed date
    reviewed_date_val = doc.get("reviewed_date")
    adopt_rows = [[Paragraph("Date of Adoption", bold_s), Paragraph(adoption, centered)]]
    if reviewed_date_val:
        adopt_rows.append([Paragraph("Last Reviewed", bold_s), Paragraph(str(reviewed_date_val), centered)])
    adopt_t = Table(adopt_rows, colWidths=[5*cm, 5*cm])
    adopt_t.setStyle(TableStyle([
        ("GRID",       (0,0),(-1,-1), 0.5, colors.HexColor("#999999")),
        ("BACKGROUND", (0,0),(0,-1),  LIGHT_BLUE),
        ("FONTSIZE",   (0,0),(-1,-1), 8),
        ("VALIGN",     (0,0),(-1,-1), "MIDDLE"),
        ("ALIGN",      (1,0),(-1,-1), "CENTER"),
    ]))
    story.append(adopt_t)
    story.append(Spacer(1, 0.6*cm))

    # Revision history
    story.append(Paragraph("REVISION HISTORY", h1s))
    if revisions:
        rh_data = [[
            Paragraph("Revision", centered_b),
            Paragraph("Date", centered_b),
            Paragraph("Status", centered_b),
            Paragraph("Description", centered_b),
        ]]
        for r in revisions:
            rh_data.append([
                Paragraph(str(r.get("revision","")).zfill(2), centered),
                Paragraph(str(r.get("revised_date","")), centered),
                Paragraph(r.get("status",""), centered),
                Paragraph(r.get("description",""), normal),
            ])
        rh_t = Table(rh_data, colWidths=[2*cm, 3*cm, 3.5*cm, None])
        rh_t.setStyle(TableStyle([
            ("GRID",            (0,0),(-1,-1), 0.5, colors.HexColor("#999999")),
            ("BACKGROUND",      (0,0),(-1,0),  NAPCO_BLUE),
            ("TEXTCOLOR",       (0,0),(-1,0),  colors.white),
            ("FONTSIZE",        (0,0),(-1,-1), 8),
            ("VALIGN",          (0,0),(-1,-1), "MIDDLE"),
            ("ROWBACKGROUNDS",  (0,1),(-1,-1), [colors.white, LIGHT_BLUE]),
        ]))
        story.append(rh_t)

    story.append(PageBreak())

    # ── TABLE OF CONTENTS ───────────────────────────────────
    story.append(Paragraph("Table of Contents", h1s))
    story.append(Spacer(1, 0.3*cm))

    toc_entries = [
        ("1.0", "Introduction", False),
        ("1.1", "Purpose", True),
        ("1.2", "Policy", True),
        ("1.3", "Scope of Application", True),
        ("1.4", "Authorities & Responsibilities", True),
        ("2.0", "Abbreviations and Definitions", False),
        ("3.0", "Procedure (Narrative or Flowchart)", False),
        ("4.0", "Associated Documentation", False),
        ("4.1", "Related Documents", True),
        ("4.2", "Resulting Records", True),
        ("4.3", "Internal / External References", True),
    ]

    toc_num_main = ParagraphStyle("TNM", fontName="Helvetica-Bold", fontSize=10, leading=14)
    toc_num_sub  = ParagraphStyle("TNS", fontName="Helvetica",      fontSize=9,  leading=13, leftIndent=10)
    toc_lbl_main = ParagraphStyle("TLM", fontName="Helvetica-Bold", fontSize=10, leading=14)
    toc_lbl_sub  = ParagraphStyle("TLS", fontName="Helvetica",      fontSize=9,  leading=13)
    toc_dots_sty = ParagraphStyle("TD",  fontName="Helvetica",      fontSize=9,  leading=13,
                                   textColor=colors.HexColor("#BBBBBB"))

    toc_rows = []
    for num, lbl, is_sub in toc_entries:
        num_sty = toc_num_sub if is_sub else toc_num_main
        lbl_sty = toc_lbl_sub if is_sub else toc_lbl_main
        toc_rows.append([
            Paragraph(num, num_sty),
            Paragraph(lbl, lbl_sty),
            Paragraph("." * 55, toc_dots_sty),
        ])

    # Number column widened to comfortably fit "1.4" / "4.3" etc. in bold
    # 10pt without being forced to wrap character-by-character.
    toc_t = Table(toc_rows, colWidths=[1.8*cm, 9*cm, None])
    toc_t.setStyle(TableStyle([
        ("VALIGN",        (0,0), (-1,-1), "BOTTOM"),
        ("BOTTOMPADDING", (0,0), (-1,-1), 4),
        ("TOPPADDING",    (0,0), (-1,-1), 0),
        ("LEFTPADDING",   (0,0), (-1,-1), 0),
        ("RIGHTPADDING",  (0,0), (-1,-1), 4),
    ]))
    story.append(toc_t)

    story.append(PageBreak())

    # ── 1.0 INTRODUCTION ────────────────────────────────────
    story.append(Paragraph("1.0 Introduction", h1s))
    story.append(HRFlowable(width="100%", thickness=0.5, color=NAPCO_BLUE))
    story.append(Spacer(1, 0.2*cm))

    purpose = doc.get("purpose","")
    if purpose:
        story.append(Paragraph("1.1 Purpose", h2s))
        story.append(Paragraph(f"1.1.1 &nbsp;&nbsp; {purpose}", numbered))
        story.append(Spacer(1, 0.3*cm))

    policy = doc.get("policy") or []
    if policy:
        story.append(Paragraph("1.2 Policy", h2s))
        for i, p in enumerate(policy, 1):
            story.append(Paragraph(f"1.2.{i} &nbsp;&nbsp; {p}", numbered))
        story.append(Spacer(1, 0.3*cm))

    scope = doc.get("scope") or []
    if scope:
        story.append(Paragraph("1.3 Scope of Application", h2s))
        for i, s in enumerate(scope, 1):
            story.append(Paragraph(f"1.3.{i} &nbsp;&nbsp; {s}", numbered))
        story.append(Spacer(1, 0.3*cm))

    resp = doc.get("responsibilities") or []
    if resp:
        story.append(Paragraph("1.4 Authorities & Responsibilities", h2s))
        for i, r in enumerate(resp, 1):
            story.append(Paragraph(f"1.4.{i} &nbsp;&nbsp; {r}", numbered))
        story.append(Spacer(1, 0.3*cm))

    # ── 2.0 ABBREVIATIONS ───────────────────────────────────
    story.append(HRFlowable(width="100%", thickness=0.5, color=NAPCO_BLUE))
    story.append(Paragraph("2.0 Abbreviations and Definitions", h1s))
    abbrevs = doc.get("abbreviations") or []
    if abbrevs:
        ab_data = [[Paragraph("Terms & Abbreviations", centered_b),
                    Paragraph("Definition", centered_b)]]
        for i, ab in enumerate(abbrevs):
            ab_data.append([
                Paragraph(ab.get("term",""), bold_s),
                Paragraph(ab.get("definition",""), normal),
            ])
        ab_t = Table(ab_data, colWidths=[5*cm, None])
        ab_t.setStyle(TableStyle([
            ("GRID",           (0,0),(-1,-1), 0.5, colors.HexColor("#999999")),
            ("BACKGROUND",     (0,0),(-1,0),  NAPCO_BLUE),
            ("TEXTCOLOR",      (0,0),(-1,0),  colors.white),
            ("FONTSIZE",       (0,0),(-1,-1), 8),
            ("VALIGN",         (0,0),(-1,-1), "MIDDLE"),
            ("ROWBACKGROUNDS", (0,1),(-1,-1), [colors.white, LIGHT_BLUE]),
        ]))
        story.append(ab_t)
    story.append(Spacer(1, 0.4*cm))

    # ── 3.0 PROCEDURE ───────────────────────────────────────
    story.append(HRFlowable(width="100%", thickness=0.5, color=NAPCO_BLUE))
    story.append(Paragraph("3.0 Procedure (Narrative or Flowchart)", h1s))

    steps = doc.get("procedure_steps") or []
    step_num = 0
    for step in steps:
        stype = step.get("type", "step")

        if stype == "input_row":
            items = step.get("items", [])
            titles = "; ".join(it.get("title", "") for it in items)
            heading = step.get("heading", "Trigger Sources")
            intro   = step.get("intro", "This process may be triggered by any of the following")
            story.append(Paragraph(heading, h2s))
            story.append(Paragraph(f"{intro}: {titles}.", numbered))
            story.append(Spacer(1, 0.15*cm))
            continue

        if stype == "branch_split":
            for br in step.get("branches", []):
                story.append(Paragraph(f"<b>{br.get('label','')}</b>", numbered))
                for s in br.get("steps", []):
                    story.append(Paragraph(f"• {s.get('title','')}", numbered))
            story.append(Spacer(1, 0.15*cm))
            continue

        step_num += 1
        story.append(Paragraph(f"{step_num}. {step.get('title','')}", h2s))
        if step.get("text"):
            story.append(Paragraph(step["text"], numbered))
        if step.get("side_branch"):
            sb_title = step["side_branch"].get("title", "")
            story.append(Paragraph(f"→ In parallel: {sb_title}", numbered))
        story.append(Spacer(1, 0.15*cm))

    # Process flowchart — paginated so it never exceeds one frame's height.
    # Box heights are dynamic (based on wrapped title length), so chunks are
    # built greedily by actually measuring each candidate chart's height.
    # A document with a small, consistently-reused set of roles and no
    # branch/merge points renders as a cross-functional swimlane grid
    # (much more readable); anything with branch_split/input_row keeps the
    # single-column renderer since lane-grid can't express branching.
    if steps:
        story.append(Spacer(1, 0.4*cm))
        story.append(Paragraph("Process Flowchart", h2s))
        story.append(Spacer(1, 0.2*cm))

        use_lane_grid = _should_use_lane_grid(steps)
        if use_lane_grid:
            ChartClass = LaneGridFlowchart
            lane_order = []
            for s in steps:
                lane = (s.get("swimlane") or "General").strip() or "General"
                if lane not in lane_order:
                    lane_order.append(lane)
            chart_kwargs = {"lane_order": lane_order}
        else:
            ChartClass = SwimlaneFlowchart
            chart_kwargs = {}

        avail_h = (PAGE_H - TOP_M - BOT_M) - 1*cm  # safety margin
        chunks, current = [], []
        for step in steps:
            trial = current + [step]
            if ChartClass(trial, avail_w, **chart_kwargs).total_h > avail_h and current:
                chunks.append(current)
                current = [step]
            else:
                current = trial
        if current:
            chunks.append(current)

        for i, chunk in enumerate(chunks):
            chart = ChartClass(chunk, avail_w, **chart_kwargs)
            story.append(chart)
            if i < len(chunks) - 1:
                story.append(PageBreak())
                story.append(Paragraph("Process Flowchart (continued)", h2s))
                story.append(Spacer(1, 0.2*cm))
        story.append(Spacer(1, 0.4*cm))

    # ── 4.0 ASSOCIATED DOCUMENTATION ────────────────────────
    story.append(HRFlowable(width="100%", thickness=0.5, color=NAPCO_BLUE))
    story.append(Paragraph("4.0 Associated Documentation", h1s))

    def ref_table(items, label):
        if not items: return
        story.append(Paragraph(label, h2s))
        rd = [[Paragraph("Code", centered_b), Paragraph("Title", centered_b)]]
        for r in items:
            rd.append([Paragraph(r.get("code","—"), normal),
                       Paragraph(r.get("title","—"), normal)])
        rt = Table(rd, colWidths=[5*cm, None])
        rt.setStyle(TableStyle([
            ("GRID",           (0,0),(-1,-1), 0.5, colors.HexColor("#999999")),
            ("BACKGROUND",     (0,0),(-1,0),  NAPCO_BLUE),
            ("TEXTCOLOR",      (0,0),(-1,0),  colors.white),
            ("FONTSIZE",       (0,0),(-1,-1), 8),
            ("VALIGN",         (0,0),(-1,-1), "MIDDLE"),
            ("ROWBACKGROUNDS", (0,1),(-1,-1), [colors.white, LIGHT_BLUE]),
        ]))
        story.append(rt)
        story.append(Spacer(1, 0.3*cm))

    ref_table(doc.get("related_docs") or [],     "4.1 Related Documents")
    ref_table(doc.get("resulting_records") or [], "4.2 Resulting Records")
    ref_table(doc.get("ext_references") or [],    "4.3 Internal / External References")

    pdf.build(story, onFirstPage=header_footer, onLaterPages=header_footer)
    buffer.seek(0)
    return buffer.read()


# ─────────────────────────────────────────────
# WORD (DOCX) GENERATION — mirrors generate_pdf's structure and content,
# including the same flowchart visuals, rasterized to images.
# ─────────────────────────────────────────────
def _flowable_to_png_bytes(flowable, width_pt, height_pt, dpi=150):
    """Renders a ReportLab Flowable (e.g. SwimlaneFlowchart/LaneGridFlowchart)
    to PNG bytes by drawing it onto a single-page PDF sized exactly to its
    own dimensions, then rasterizing with pdftoppm (poppler-utils)."""
    from reportlab.pdfgen import canvas as rl_canvas
    buf = io.BytesIO()
    c = rl_canvas.Canvas(buf, pagesize=(width_pt, height_pt))
    flowable.drawOn(c, 0, 0)
    c.showPage()
    c.save()
    pdf_bytes = buf.getvalue()

    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tf:
        tf.write(pdf_bytes)
        tf_path = tf.name
    out_prefix = tf_path[:-4]
    try:
        subprocess.run(["pdftoppm", "-png", "-r", str(dpi), tf_path, out_prefix],
                        check=True, capture_output=True)
        png_path = out_prefix + "-1.png"
        with open(png_path, "rb") as f:
            png_bytes = f.read()
    finally:
        for p in (tf_path, out_prefix + "-1.png"):
            if os.path.exists(p):
                os.remove(p)
    return png_bytes


def _docx_set_cell_background(cell, hex_color):
    """python-docx has no direct API for cell shading — patches the cell's
    underlying XML directly."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def _docx_set_table_borders(table, size=4, color="999999"):
    """Adds a simple grid border to every cell of the table. Relying on a
    named 'Table Grid' style can silently fail depending on the base
    template, so this sets the border XML directly for guaranteed results."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tblPr.append(borders)


def _docx_apply_font_everywhere(document, font_name="Trebuchet MS"):
    """Sets the default document font AND force-applies it to every run
    already created (paragraphs, headings, tables, header/footer) — relying
    on style cascade alone can render inconsistently across Word/LibreOffice."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    document.styles["Normal"].font.name = font_name
    for style_name in ["Heading 1", "Heading 2", "Heading 3", "Heading 4"]:
        try:
            document.styles[style_name].font.name = font_name
        except KeyError:
            pass

    def _set_run_font(run):
        run.font.name = font_name
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:ascii"), font_name)
        rFonts.set(qn("w:hAnsi"), font_name)
        rFonts.set(qn("w:cs"), font_name)

    def _walk_paragraphs(paragraphs):
        for p in paragraphs:
            for r in p.runs:
                _set_run_font(r)

    def _walk_tables(tables):
        for t in tables:
            for row in t.rows:
                for cell in row.cells:
                    _walk_paragraphs(cell.paragraphs)

    _walk_paragraphs(document.paragraphs)
    _walk_tables(document.tables)
    for section in document.sections:
        _walk_paragraphs(section.header.paragraphs)
        _walk_tables(section.header.tables)
        _walk_paragraphs(section.footer.paragraphs)
        _walk_tables(section.footer.tables)


def _docx_add_footer_disclaimer(document):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    section = document.sections[0]
    footer = section.footer
    p1 = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p1.text = ("THE INFORMATION CONTAINED HEREIN IS PROPRIETARY TO NAPCO NATIONAL AND IT SHALL NOT BE "
               "USED, REPRODUCED OR DISCLOSED TO OTHERS EXCEPT AS SPECIFICALLY PERMITTED IN WRITING BY "
               "THE PROPRIETOR.")
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p1.runs:
        run.font.size = Pt(6.5)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    p2 = footer.add_paragraph('"UNCONTROLLED IF PRINTED"')
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p2.runs:
        run.font.size = Pt(7)
        run.font.bold = True
        run.font.italic = True


def _docx_add_header_table(document, doc_code, title_txt, adoption, issue_date, rev_date):
    """Adds the running header (logo + control-number table) to the section."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    section = document.sections[0]
    header = section.header

    try:
        p_logo = header.paragraphs[0]
        run = p_logo.add_run()
        run.add_picture(LOGO_PATH, height=Cm(1.0))
    except Exception:
        pass

    tbl = header.add_table(rows=2, cols=6, width=Cm(17))
    tbl.autofit = False
    _docx_set_table_borders(tbl)
    widths = [Cm(2.2), Cm(7.6), Cm(2.8), Cm(2.2), Cm(2.2), Cm(2.0)]
    hdr_cells = tbl.rows[0].cells
    hdr_cells[0].merge(hdr_cells[0])
    row1 = ["TITLE", title_txt.upper(), "", "", "", ""]
    # Row 1: TITLE | value (merged) | (blank) | (blank) | NBR OF PAGES-ish | (blank)
    r1 = tbl.rows[0].cells
    r1[0].text = "TITLE"
    r1[1].merge(r1[4])
    r1[1].text = title_txt.upper()
    r1[5].text = "QUALITY PROCEDURE"

    r2 = tbl.rows[1].cells
    r2[0].text = "CONTROL NBR."
    r2[1].text = str(doc_code)
    r2[2].text = "1ST ISSUE DATE"
    r2[3].text = str(issue_date)
    r2[4].text = "REVISION DATE"
    r2[5].text = str(rev_date)

    for ci, w in enumerate(widths):
        for row in tbl.rows:
            row.cells[ci].width = w

    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(7)
            _docx_set_cell_background(cell, "F2F2F2")
    for cell in [r1[0], r2[0], r2[2], r2[4]]:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True


def generate_docx(doc):
    """Generates a Word (.docx) version of the document, visually matching
    generate_pdf — same sections, same flowchart images (rasterized from the
    identical drawing code used for the PDF), same tables."""
    NAPCO_BLUE_HEX = "0D68A3"
    LIGHT_BLUE_HEX = "D5E8F0"

    document = Document()
    section = document.sections[0]
    section.page_width  = Cm(21.0)
    section.page_height  = Cm(29.7)
    section.left_margin  = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin   = Cm(2.6)
    section.bottom_margin = Cm(1.8)

    doc_code   = doc.get("doc_code", "-")
    title_txt  = doc.get("title", "")
    dept_label = doc.get("dept_label") or doc.get("dept", "")
    adoption   = doc.get("date_of_adoption", "-")
    revisions  = doc.get("_revisions") or []
    issue_date = revisions[0]["revised_date"] if revisions else adoption
    rev_date   = revisions[-1]["revised_date"] if revisions else adoption

    _docx_add_header_table(document, doc_code, title_txt, adoption, issue_date, rev_date)
    _docx_add_footer_disclaimer(document)

    avail_w_pt = (21.0 - 4.0) * cm  # content width in points, matching PDF's avail_w

    # ── Cover page ──────────────────────────────
    for _ in range(3):
        document.add_paragraph()
    try:
        p_logo = document.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_logo.add_run().add_picture(LOGO_PATH, height=Cm(2.2))
    except Exception:
        pass

    p_dept = document.add_paragraph()
    p_dept.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_dept.add_run(dept_label.upper())
    r.font.size = Pt(13)
    r.font.bold = True

    p_title = document.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_title.add_run(title_txt.upper())
    r.font.size = Pt(22)
    r.font.bold = True
    document.add_paragraph()

    approvals = doc.get("approvals") or []
    if approvals:
        n = len(approvals)
        ap = document.add_table(rows=5, cols=n + 1)
        ap.alignment = WD_TABLE_ALIGNMENT.CENTER
        _docx_set_table_borders(ap)
        labels = ["APPROVED BY", "Department", "Function", "Signature", "Date"]
        for ri, lbl in enumerate(labels):
            ap.rows[ri].cells[0].text = lbl
            for c in ap.rows[ri].cells:
                for p in c.paragraphs:
                    for run in p.runs:
                        run.font.bold = (ri == 0 or True) if False else run.font.bold
        for ci, a in enumerate(approvals, start=1):
            ap.rows[1].cells[ci].text = a.get("department", "")
            ap.rows[2].cells[ci].text = a.get("function", "")
        for ri in range(5):
            _docx_set_cell_background(ap.rows[ri].cells[0], LIGHT_BLUE_HEX)
        _docx_set_cell_background(ap.rows[0].cells[0], NAPCO_BLUE_HEX)
        for c in ap.rows[0].cells:
            _docx_set_cell_background(c, NAPCO_BLUE_HEX)
            for p in c.paragraphs:
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.bold = True
        for row in ap.rows:
            for c in row.cells:
                for p in c.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.font.size = Pt(8)
        document.add_paragraph()

    adopt_rows = 2 if doc.get("reviewed_date") else 1
    adopt_t = document.add_table(rows=adopt_rows, cols=2)
    adopt_t.alignment = WD_TABLE_ALIGNMENT.CENTER
    _docx_set_table_borders(adopt_t)
    adopt_t.rows[0].cells[0].text = "Date of Adoption"
    adopt_t.rows[0].cells[1].text = str(adoption)
    if doc.get("reviewed_date"):
        adopt_t.rows[1].cells[0].text = "Last Reviewed"
        adopt_t.rows[1].cells[1].text = str(doc.get("reviewed_date"))
    for row in adopt_t.rows:
        _docx_set_cell_background(row.cells[0], LIGHT_BLUE_HEX)
        for c in row.cells:
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(9)
    document.add_paragraph()

    h = document.add_heading("REVISION HISTORY", level=2)
    if revisions:
        rt = document.add_table(rows=1, cols=4)
        rt.alignment = WD_TABLE_ALIGNMENT.CENTER
        _docx_set_table_borders(rt)
        hdr = rt.rows[0].cells
        for i, lbl in enumerate(["Revision", "Date", "Status", "Description"]):
            hdr[i].text = lbl
            _docx_set_cell_background(hdr[i], NAPCO_BLUE_HEX)
            for p in hdr[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.bold = True
                    run.font.size = Pt(9)
        for rv in revisions:
            row = rt.add_row().cells
            row[0].text = f"{rv.get('revision', 0):02d}"
            row[1].text = str(rv.get("revised_date", ""))
            row[2].text = str(rv.get("status", ""))
            row[3].text = str(rv.get("description", ""))
            for c in row:
                for p in c.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(8.5)

    document.add_page_break()

    # ── Table of Contents (static, matches PDF layout) ──────────────
    document.add_heading("Table of Contents", level=1)
    toc_entries = [
        ("1.0", "Introduction", False), ("1.1", "Purpose", True), ("1.2", "Policy", True),
        ("1.3", "Scope of Application", True), ("1.4", "Authorities & Responsibilities", True),
        ("2.0", "Abbreviations and Definitions", False),
        ("3.0", "Procedure (Narrative or Flowchart)", False),
        ("4.0", "Associated Documentation", False), ("4.1", "Related Documents", True),
        ("4.2", "Resulting Records", True), ("4.3", "Internal / External References", True),
    ]
    for num, lbl, is_sub in toc_entries:
        p = document.add_paragraph()
        if is_sub:
            p.paragraph_format.left_indent = Cm(0.8)
        run = p.add_run(f"{num}\t{lbl}")
        run.font.bold = not is_sub
        run.font.size = Pt(10 if not is_sub else 9.5)
    document.add_page_break()

    # ── 1.0 Introduction ──────────────────────────────
    document.add_heading("1.0 Introduction", level=1)
    document.add_heading("1.1 Purpose", level=2)
    document.add_paragraph(doc.get("purpose", ""))

    policy = doc.get("policy") or []
    if policy:
        document.add_heading("1.2 Policy", level=2)
        for p_item in policy:
            document.add_paragraph(p_item, style="List Bullet")

    scope = doc.get("scope") or []
    if scope:
        document.add_heading("1.3 Scope of Application", level=2)
        for s_item in scope:
            document.add_paragraph(s_item, style="List Bullet")

    resp = doc.get("responsibilities") or []
    if resp:
        document.add_heading("1.4 Authorities & Responsibilities", level=2)
        for r_item in resp:
            document.add_paragraph(r_item, style="List Bullet")

    # ── 2.0 Abbreviations ──────────────────────────────
    abbrevs = doc.get("abbreviations") or []
    if abbrevs:
        document.add_heading("2.0 Abbreviations and Definitions", level=1)
        at = document.add_table(rows=1, cols=2)
        at.alignment = WD_TABLE_ALIGNMENT.CENTER
        _docx_set_table_borders(at)
        hdr = at.rows[0].cells
        hdr[0].text, hdr[1].text = "Terms & Abbreviations", "Definition"
        for c in hdr:
            _docx_set_cell_background(c, NAPCO_BLUE_HEX)
            for p in c.paragraphs:
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.bold = True
        for a in abbrevs:
            row = at.add_row().cells
            row[0].text = a.get("term", "")
            row[1].text = a.get("definition", "")
            for run in row[0].paragraphs[0].runs:
                run.font.bold = True

    # ── 3.0 Procedure ──────────────────────────────
    document.add_heading("3.0 Procedure (Narrative or Flowchart)", level=1)
    steps = doc.get("procedure_steps") or []
    step_num = 0
    for step in steps:
        stype = step.get("type", "step")
        if stype == "input_row":
            items = step.get("items", [])
            titles = "; ".join(it.get("title", "") for it in items)
            heading = step.get("heading", "Trigger Sources")
            intro   = step.get("intro", "This process may be triggered by any of the following")
            document.add_heading(heading, level=3)
            document.add_paragraph(f"{intro}: {titles}.")
            continue
        if stype == "branch_split":
            for br in step.get("branches", []):
                bp = document.add_paragraph()
                bp.add_run(br.get("label", "")).bold = True
                for s in br.get("steps", []):
                    document.add_paragraph(s.get("title", ""), style="List Bullet")
            continue
        step_num += 1
        document.add_heading(f"{step_num}. {step.get('title','')}", level=3)
        if step.get("text"):
            document.add_paragraph(step["text"])
        if step.get("side_branch"):
            sp = document.add_paragraph()
            sp.add_run(f"→ In parallel: {step['side_branch'].get('title','')}").italic = True

    # ── Flowchart image(s) ──────────────────────────────
    if steps:
        document.add_heading("Process Flowchart", level=2)
        use_lane_grid = _should_use_lane_grid(steps)
        if use_lane_grid:
            ChartClass = LaneGridFlowchart
            lane_order = []
            for s in steps:
                lane = (s.get("swimlane") or "General").strip() or "General"
                if lane not in lane_order:
                    lane_order.append(lane)
            chart_kwargs = {"lane_order": lane_order}
        else:
            ChartClass = SwimlaneFlowchart
            chart_kwargs = {}

        max_h_pt = 24 * cm  # comfortable single-page image height budget
        chunks, current = [], []
        for step in steps:
            trial = current + [step]
            if ChartClass(trial, avail_w_pt, **chart_kwargs).total_h > max_h_pt and current:
                chunks.append(current)
                current = [step]
            else:
                current = trial
        if current:
            chunks.append(current)

        for i, chunk in enumerate(chunks):
            chart = ChartClass(chunk, avail_w_pt, **chart_kwargs)
            try:
                png_bytes = _flowable_to_png_bytes(chart, chart.width, chart.total_h)
                p_img = document.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img.add_run().add_picture(io.BytesIO(png_bytes), width=Cm(17))
            except Exception as e:
                document.add_paragraph(f"[Flowchart image could not be generated: {e}]")
            if i < len(chunks) - 1:
                document.add_page_break()
                document.add_heading("Process Flowchart (continued)", level=2)

    # ── 4.0 Associated Documentation ──────────────────────────────
    document.add_heading("4.0 Associated Documentation", level=1)

    def _ref_table(refs, heading):
        if not refs:
            return
        document.add_heading(heading, level=2)
        rt = document.add_table(rows=1, cols=2)
        rt.alignment = WD_TABLE_ALIGNMENT.CENTER
        _docx_set_table_borders(rt)
        hdr = rt.rows[0].cells
        hdr[0].text, hdr[1].text = "Code", "Title"
        for c in hdr:
            _docx_set_cell_background(c, NAPCO_BLUE_HEX)
            for p in c.paragraphs:
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.bold = True
        for r in refs:
            row = rt.add_row().cells
            row[0].text = r.get("code", "")
            row[1].text = r.get("title", "")

    _ref_table(doc.get("related_docs") or [],      "4.1 Related Documents")
    _ref_table(doc.get("resulting_records") or [], "4.2 Resulting Records")
    _ref_table(doc.get("ext_references") or [],    "4.3 Internal / External References")

    _docx_apply_font_everywhere(document, "Trebuchet MS")

    buf = io.BytesIO()
    document.save(buf)
    buf.seek(0)
    return buf.read()


# ─────────────────────────────────────────────
# MAIN PAGE
# ─────────────────────────────────────────────
def show():
    require_auth()
    sb      = get_supabase()
    profile = get_profile()
    uid     = profile["id"]

    st.title("📋 Process & Procedure Builder")
    st.caption("Create, view and manage Easternpak procedures and processes in Napco NFP format")

    tab_list, tab_new, tab_master = st.tabs([
        "📂 Document Library",
        "➕ New Document",
        "📚 Master Document List",
    ])

    # ══════════════════════════════════════════
    # TAB 1 — DOCUMENT LIBRARY
    # ══════════════════════════════════════════
    with tab_list:
        col1, col2, col3 = st.columns(3)
        with col1:
            dept_f = st.selectbox("Department", ["All","SC","PL","TQA","PM"])
        with col2:
            type_f = st.selectbox("Type", ["All","PD","PR"])
        with col3:
            search = st.text_input("🔍 Search title or code")

        q = sb.table("proc_documents").select("*").order("doc_code")
        if dept_f != "All":
            q = q.eq("dept", dept_f)
        if type_f != "All":
            q = q.eq("doc_type", type_f)
        res = q.execute()
        docs = res.data or []

        if search:
            s = search.lower()
            docs = [d for d in docs if s in (d.get("title","") or "").lower()
                    or s in (d.get("doc_code","") or "").lower()]

        if not docs:
            st.info("No documents found. Create your first one using the ➕ New Document tab.")
        else:
            for doc in docs:
                rl = revision_label(doc.get("revision", 0))
                with st.expander(
                    f"**{doc.get('doc_code','—')}** — {doc.get('title','')}  "
                    f"| Rev {doc.get('revision',0):02d} | {rl}"
                ):
                    rev_res = sb.table("proc_revisions").select("*")\
                        .eq("doc_id", doc["id"]).order("revision").execute()
                    doc["_revisions"] = rev_res.data or []

                    col_dl1, col_dl2, _ = st.columns([2, 2, 6])
                    with col_dl1:
                        if st.button("⬇️ Generate PDF", key=f"gen_{doc['id']}"):
                            with st.spinner("Generating PDF…"):
                                try:
                                    pdf_bytes = generate_pdf(doc)
                                    st.download_button(
                                        label="📥 Click to Download",
                                        data=pdf_bytes,
                                        file_name=f"{doc.get('doc_code','document')}.pdf",
                                        mime="application/pdf",
                                        key=f"dl_{doc['id']}",
                                    )
                                except Exception as e:
                                    st.error(f"PDF error: {e}")
                    with col_dl2:
                        if st.button("📝 Generate Word", key=f"genw_{doc['id']}"):
                            with st.spinner("Generating Word document…"):
                                try:
                                    docx_bytes = generate_docx(doc)
                                    st.download_button(
                                        label="📥 Click to Download",
                                        data=docx_bytes,
                                        file_name=f"{doc.get('doc_code','document')}.docx",
                                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                        key=f"dlw_{doc['id']}",
                                    )
                                except Exception as e:
                                    st.error(f"Word document error: {e}")

                    _render_document(sb, doc, uid)

    # ══════════════════════════════════════════
    # TAB 2 — NEW DOCUMENT
    # ══════════════════════════════════════════
    with tab_new:
        if not can_write():
            st.info("View-only access. Contact your administrator.")
            return

        st.markdown("### Document Identity")
        c1, c2, c3 = st.columns(3)
        with c1:
            dept_sel  = st.selectbox("Department *", list(DEPT_MAP.keys()), key="nd_dept")
            dept_info = DEPT_MAP[dept_sel]
            dept_code = dept_info["code"]
            if dept_code is None:
                dept_code = st.text_input("Department code", key="nd_dept_other").upper()
        with c2:
            sub_opts    = dept_info["subs"]
            subdept_sel = st.selectbox("Sub-department *", sub_opts, key="nd_subdept")
            if subdept_sel == "Other":
                subdept_sel = st.text_input("Sub-dept code", key="nd_subdept_other").upper()
        with c3:
            type_sel = st.selectbox("Document Type *", list(DOC_TYPE_MAP.keys()), key="nd_type")
            doc_type = DOC_TYPE_MAP[type_sel]

        if dept_code and subdept_sel and doc_type:
            seq_res = sb.table("proc_documents")\
                .select("seq_number")\
                .eq("dept", dept_code).eq("subdept", subdept_sel)\
                .eq("doc_type", doc_type)\
                .order("seq_number", desc=True).limit(1).execute()
            next_seq     = (seq_res.data[0]["seq_number"] + 1) if seq_res.data else 1
            preview_code = f"NFP-EP-{dept_code}-{subdept_sel}-{doc_type}-{next_seq:02d}-00"
            st.info(f"📄 Auto-generated code: **{preview_code}**")

        st.markdown("---")
        st.markdown("### Cover Page")
        c1, c2 = st.columns(2)
        with c1:
            title      = st.text_input("Document Title *",
                                        placeholder="e.g. Customer Complaint Handling Process")
            dept_label = st.text_input("Department Name (for cover)",
                                        placeholder="e.g. Supply Chain Department")
        with c2:
            adoption_dt = st.date_input("Date of Adoption", value=date.today())

        st.markdown("**Approvals**")
        if "approvals" not in st.session_state:
            st.session_state.approvals = [
                {"department": "", "function": ""},
                {"department": "", "function": ""},
                {"department": "Top Management", "function": "Operations Manager"},
            ]
        ap_cols = st.columns(len(st.session_state.approvals))
        for i, ap in enumerate(st.session_state.approvals):
            with ap_cols[i]:
                st.session_state.approvals[i]["department"] = st.text_input(
                    f"Dept {i+1}", value=ap["department"], key=f"ap_dept_{i}")
                st.session_state.approvals[i]["function"] = st.text_input(
                    f"Function {i+1}", value=ap["function"], key=f"ap_func_{i}")
        if st.button("➕ Add approver"):
            st.session_state.approvals.append({"department":"","function":""})
            st.rerun()

        st.markdown("---")
        st.markdown("### 1.0 Introduction")
        purpose = st.text_area("1.1 Purpose", height=100,
                                placeholder="The purpose of this procedure is to…")

        st.markdown("**1.2 Policy**")
        if "policy_points" not in st.session_state:
            st.session_state.policy_points = [""]
        _edit_list("policy_points", "Policy point")

        st.markdown("**1.3 Scope of Application**")
        if "scope_points" not in st.session_state:
            st.session_state.scope_points = [""]
        _edit_list("scope_points", "Scope point")

        st.markdown("**1.4 Authorities & Responsibilities**")
        if "resp_points" not in st.session_state:
            st.session_state.resp_points = [""]
        _edit_list("resp_points", "Responsibility point")

        st.markdown("---")
        st.markdown("### 2.0 Abbreviations & Definitions")
        if "abbrevs" not in st.session_state:
            st.session_state.abbrevs = [{"term":"","definition":""}]
        for i, ab in enumerate(st.session_state.abbrevs):
            c1, c2, c3 = st.columns([2,4,1])
            with c1:
                st.session_state.abbrevs[i]["term"] = st.text_input(
                    "Term", value=ab["term"], key=f"ab_term_{i}",
                    label_visibility="collapsed", placeholder="Term")
            with c2:
                st.session_state.abbrevs[i]["definition"] = st.text_input(
                    "Def", value=ab["definition"], key=f"ab_def_{i}",
                    label_visibility="collapsed", placeholder="Definition")
            with c3:
                if st.button("🗑️", key=f"ab_del_{i}") and len(st.session_state.abbrevs) > 1:
                    st.session_state.abbrevs.pop(i)
                    st.rerun()
        if st.button("➕ Add abbreviation"):
            st.session_state.abbrevs.append({"term":"","definition":""})
            st.rerun()

        st.markdown("---")
        st.markdown("### 3.0 Procedure Steps & Flowchart")
        with st.expander("⚙️ Configure swimlanes"):
            lanes_raw = st.text_area("Swimlanes (one per line)",
                                     value="\n".join(SWIMLANES_DEFAULT),
                                     height=150, key="nd_lanes")
            lanes = [l.strip() for l in lanes_raw.split("\n") if l.strip()]

        if "proc_steps" not in st.session_state:
            st.session_state.proc_steps = []

        for i, step in enumerate(st.session_state.proc_steps):
            with st.expander(f"Step {i+1}: {step.get('title','Untitled')}", expanded=False):
                c1, c2, c3 = st.columns([3,2,1])
                with c1:
                    st.session_state.proc_steps[i]["title"] = st.text_input(
                        "Step title", value=step.get("title",""), key=f"ps_title_{i}")
                with c2:
                    idx = lanes.index(step.get("swimlane", lanes[0])) \
                          if step.get("swimlane") in lanes else 0
                    st.session_state.proc_steps[i]["swimlane"] = st.selectbox(
                        "Swimlane", lanes, key=f"ps_lane_{i}", index=idx)
                with c3:
                    shapes = ["rect","diamond","rounded"]
                    sidx   = shapes.index(step.get("shape","rect")) \
                             if step.get("shape") in shapes else 0
                    st.session_state.proc_steps[i]["shape"] = st.selectbox(
                        "Shape", shapes, key=f"ps_shape_{i}", index=sidx)
                st.session_state.proc_steps[i]["text"] = st.text_area(
                    "Description", value=step.get("text",""),
                    key=f"ps_text_{i}", height=80)
                st.session_state.proc_steps[i]["connection_label"] = st.text_input(
                    "Arrow label to next step",
                    value=step.get("connection_label",""),
                    key=f"ps_conn_{i}", placeholder="e.g. Yes / No")
                if st.button("🗑️ Remove step", key=f"ps_del_{i}"):
                    st.session_state.proc_steps.pop(i)
                    st.rerun()

        if st.button("➕ Add step"):
            import uuid
            st.session_state.proc_steps.append({
                "id": str(uuid.uuid4())[:8],
                "title":"", "text":"",
                "swimlane": lanes[0] if lanes else "General",
                "shape":"rect", "connection_label":"",
            })
            st.rerun()

        st.markdown("---")
        st.markdown("### 4.0 Associated Documentation")
        st.markdown("**4.1 Related Documents**")
        if "rel_docs" not in st.session_state:
            st.session_state.rel_docs = [{"code":"","title":""}]
        _edit_refs("rel_docs", sb, uid)

        st.markdown("**4.2 Resulting Records**")
        if "rec_docs" not in st.session_state:
            st.session_state.rec_docs = [{"code":"","title":""}]
        _edit_refs("rec_docs", sb, uid)

        st.markdown("**4.3 Internal / External References**")
        if "ext_refs" not in st.session_state:
            st.session_state.ext_refs = [{"code":"","title":""}]
        _edit_refs("ext_refs", sb, uid)

        st.markdown("---")
        if st.button("💾 Save Document", type="primary"):
            if not title or not dept_code or not subdept_sel or not doc_type:
                st.error("Title, department, sub-department and type are required.")
            else:
                mermaid_code = build_mermaid(st.session_state.proc_steps, lanes)
                try:
                    res = sb.table("proc_documents").insert({
                        "dept":              dept_code,
                        "subdept":           subdept_sel,
                        "doc_type":          doc_type,
                        "seq_number":        0,
                        "revision":          0,
                        "title":             title,
                        "dept_label":        dept_label or None,
                        "date_of_adoption":  adoption_dt.isoformat(),
                        "purpose":           purpose or None,
                        "policy":            [p for p in st.session_state.policy_points if p],
                        "scope":             [s for s in st.session_state.scope_points if s],
                        "responsibilities":  [r for r in st.session_state.resp_points if r],
                        "abbreviations":     [a for a in st.session_state.abbrevs if a["term"]],
                        "procedure_steps":   st.session_state.proc_steps,
                        "flowchart_mermaid": mermaid_code or None,
                        "related_docs":      [r for r in st.session_state.rel_docs
                                              if r["code"] or r["title"]],
                        "resulting_records": [r for r in st.session_state.rec_docs
                                              if r["code"] or r["title"]],
                        "ext_references":    [r for r in st.session_state.ext_refs
                                              if r["code"] or r["title"]],
                        "approvals":         st.session_state.approvals,
                        "created_by":        uid,
                        "updated_by":        uid,
                    }).execute()

                    new_doc        = res.data[0]
                    new_id         = new_doc["id"]
                    doc_code_final = new_doc.get("doc_code", preview_code)

                    all_refs = (
                        [(r,"related_doc")      for r in st.session_state.rel_docs] +
                        [(r,"resulting_record") for r in st.session_state.rec_docs] +
                        [(r,"reference")        for r in st.session_state.ext_refs]
                    )
                    for ref, rtype in all_refs:
                        code   = ref.get("code","").strip()
                        rtitle = ref.get("title","").strip()
                        if not code and not rtitle: continue
                        check_or_add_master(sb, code or rtitle, rtitle or code,
                                            "Unknown", False, uid)
                        master = sb.table("master_documents").select("id")\
                            .eq("doc_code", code or rtitle).execute()
                        mid = master.data[0]["id"] if master.data else None
                        try:
                            sb.table("doc_references").insert({
                                "source_doc_id": new_id,
                                "ref_type":      rtype,
                                "master_doc_id": mid,
                                "raw_code":      code or None,
                                "raw_title":     rtitle or None,
                            }).execute()
                        except Exception:
                            pass

                    check_or_add_master(sb, doc_code_final, title, doc_type, True, uid)
                    st.success(f"✅ Document **{doc_code_final}** saved successfully!")
                    _clear_form()
                    st.rerun()
                except Exception as e:
                    st.error(f"Error saving: {e}")

    # ══════════════════════════════════════════
    # TAB 3 — MASTER DOCUMENT LIST
    # ══════════════════════════════════════════
    with tab_master:
        st.markdown("### 📚 Master Document List")
        st.caption("All documents referenced anywhere in the system — internal and external.")

        # ── Upload an already-existing document (form/policy/work instruction) ──
        with st.expander("📤 Upload an Existing Document"):
            st.caption(
                "For forms, policies, and work instructions that are already finished and "
                "approved outside this app — upload the file and log its control details here, "
                "rather than rebuilding it through the procedure builder."
            )
            st.info(
                "**To attach a file to a document already in the master list** (e.g. an Excel "
                "master list, a policy that was added via SQL), type its **exact** Document Code "
                "and Title below — matching exactly will update that existing entry instead of "
                "creating a duplicate. Example: code `NFP-EP-PL-XX-FM-01`, "
                "title `Master List of Instruments`."
            )
            with st.form("upload_existing_doc", clear_on_submit=True):
                uc1, uc2 = st.columns(2)
                with uc1:
                    up_code  = st.text_input("Document Code*", placeholder="e.g. FM:03.10 or NFP-EP-...")
                    up_title = st.text_input("Title*", placeholder="e.g. Corrective & Preventive Action Request")
                    up_cat   = st.selectbox("Category*",
                                             ["Procedure","Policy","Work Instruction","Form","External Standard","Other"])
                    up_internal = st.checkbox("Internal document", value=True)
                with uc2:
                    up_issue_date = st.date_input("Issue Date", value=None)
                    up_rev_label  = st.text_input("Revision (as printed on the document)", placeholder="e.g. Rev 5 or 00")
                    up_rev_date   = st.date_input("Revision Date", value=None)
                    up_approved   = st.text_input("Approved By", placeholder="e.g. Executive Director / ED")

                up_reviewed_date = st.date_input(
                    "Reviewed Date (auto-set to today since uploading is how you record a review — "
                    "change this only if the document's real review date is different)",
                    value=date.today()
                )
                up_status  = st.selectbox("Status", ["Active","Draft","Obsolete","Superseded"])
                up_file    = st.file_uploader("File", type=["pdf","docx","doc","xlsx","xls","png","jpg","jpeg"])

                submitted = st.form_submit_button("Upload & Save")
                if submitted:
                    up_code_clean  = (up_code or "").strip()
                    up_title_clean = (up_title or "").strip()
                    if not up_code_clean or not up_title_clean:
                        st.error(
                            "Document Code and Title are required. "
                            f"Received code='{up_code!r}', title='{up_title!r}' — "
                            "make sure both fields are filled in before clicking Upload & Save."
                        )
                    else:
                        try:
                            file_url, file_name = None, None
                            if up_file is not None:
                                file_bytes = up_file.getvalue()
                                storage_path = f"{up_code_clean.replace('/', '-')}/{up_file.name}"
                                sb.storage.from_("documents").upload(
                                    storage_path, file_bytes,
                                    {"content-type": up_file.type or "application/octet-stream",
                                     "upsert": "true"})
                                file_url  = sb.storage.from_("documents").get_public_url(storage_path)
                                file_name = up_file.name

                            payload = {
                                "doc_code":      up_code_clean,
                                "title":         up_title_clean,
                                "doc_type":      up_cat,
                                "category":      up_cat,
                                "is_internal":   up_internal,
                                "status":        up_status,
                                "approved_by":   up_approved.strip() or None,
                                "issue_date":    up_issue_date.isoformat() if up_issue_date else None,
                                "revision_label":up_rev_label.strip() or None,
                                "revision_date": up_rev_date.isoformat() if up_rev_date else None,
                                "reviewed_date": up_reviewed_date.isoformat() if up_reviewed_date else None,
                                "created_by":    uid,
                            }
                            if file_url:
                                payload["file_url"]  = file_url
                                payload["file_name"] = file_name

                            existing = sb.table("master_documents").select("id").eq("doc_code", up_code_clean).execute()
                            if existing.data:
                                sb.table("master_documents").update(payload).eq("doc_code", up_code_clean).execute()
                                st.success(f"Updated existing master list entry for {up_code_clean}.")
                            else:
                                sb.table("master_documents").insert(payload).execute()
                                st.success(f"Added {up_code_clean} to the master list.")
                            st.rerun()
                        except Exception as e:
                            st.error(f"Error uploading document: {e}")

        res     = sb.table("master_documents").select("*").order("doc_code").execute()
        masters = res.data or []

        today = date.today()

        def _next_review(m):
            """Next review = reviewed_date + 1 year. Returns None if never reviewed."""
            rd_raw = m.get("reviewed_date")
            if not rd_raw:
                return None
            try:
                rd = date.fromisoformat(str(rd_raw)[:10])
                return rd.replace(year=rd.year + 1)
            except Exception:
                return None

        c1, c2, c3, c4 = st.columns(4)
        with c1:
            int_f = st.selectbox("Source", ["All","Internal","External"])
        with c2:
            msearch = st.text_input("Search", placeholder="code or title")
        with c3:
            review_f = st.selectbox("Review status",
                                     ["All","Never reviewed","Due soon (30 days)","Overdue","Obsolete"])
        with c4:
            cat_options = ["All","Procedure","Policy","Work Instruction","Form","External Standard","Other"]
            cat_f = st.selectbox("Category", cat_options)

        if int_f == "Internal":
            masters = [m for m in masters if m.get("is_internal")]
        elif int_f == "External":
            masters = [m for m in masters if not m.get("is_internal")]
        if msearch:
            s       = msearch.lower()
            masters = [m for m in masters
                       if s in (m.get("doc_code","") or "").lower()
                       or s in (m.get("title","") or "").lower()]
        if cat_f != "All":
            masters = [m for m in masters
                       if (m.get("category") or m.get("doc_type") or "") == cat_f]
        if review_f != "All":
            def _matches(m):
                is_obsolete = (m.get("status") == "Obsolete")
                if review_f == "Obsolete":
                    return is_obsolete
                if is_obsolete:
                    # Obsolete docs are out of the active review cycle —
                    # they don't count as never-reviewed/due/overdue.
                    return False
                nr = _next_review(m)
                if review_f == "Never reviewed":
                    return nr is None
                if nr is None:
                    return False
                if review_f == "Overdue":
                    return nr <= today
                if review_f == "Due soon (30 days)":
                    return today < nr <= today + timedelta(days=30)
                return True
            masters = [m for m in masters if _matches(m)]

        if not masters:
            st.info("No documents in master list yet.")
        else:
            st.caption(f"{len(masters)} document(s) in master list")

            hdr_cols = st.columns([1.5, 2.6, 1.1, 1, 1.3, 1.3, 0.6, 0.6, 0.6])
            for col, label in zip(hdr_cols,
                                   ["Code","Title","Category","Source","Reviewed Date","Next Review","","",""]):
                col.markdown(f"**{label}**")
            st.markdown("<hr style='margin:2px 0'>", unsafe_allow_html=True)

            for m in masters:
                mid = m["id"]
                is_obsolete = (m.get("status") == "Obsolete")
                row = st.columns([1.5, 2.6, 1.1, 1, 1.3, 1.3, 0.6, 0.6, 0.6])

                code_txt  = m.get("doc_code","—")
                title_txt = m.get("title","")
                if is_obsolete:
                    strike = "color:#999;text-decoration:line-through"
                    row[0].markdown(f"<span style='{strike}'>{code_txt}</span>", unsafe_allow_html=True)
                    row[1].markdown(f"<span style='{strike}'>{title_txt}</span>", unsafe_allow_html=True)
                else:
                    row[0].markdown(code_txt)
                    row[1].markdown(title_txt)
                row[2].markdown(m.get("category") or m.get("doc_type","—"))
                row[3].markdown("Internal" if m.get("is_internal") else "External")

                rd_raw = m.get("reviewed_date")
                try:
                    rd_val = date.fromisoformat(str(rd_raw)[:10]) if rd_raw else today
                except Exception:
                    rd_val = today
                new_reviewed = row[4].date_input("Reviewed", value=rd_val,
                                                  key=f"revdate_{mid}", label_visibility="collapsed")

                if is_obsolete:
                    row[5].markdown("⬛ *Obsolete — excluded from reviews*")
                else:
                    next_review = _next_review(m)
                    if next_review is None:
                        row[5].markdown("⚪ Never reviewed")
                    elif next_review <= today:
                        row[5].markdown(f"🔴 {next_review.isoformat()} (overdue)")
                    elif (next_review - today).days <= 30:
                        row[5].markdown(f"🟡 {next_review.isoformat()}")
                    else:
                        row[5].markdown(f"🟢 {next_review.isoformat()}")

                if row[6].button("💾", key=f"save_{mid}", help="Save reviewed date"):
                    try:
                        sb.table("master_documents").update(
                            {"reviewed_date": new_reviewed.isoformat()}
                        ).eq("id", mid).execute()
                        st.success(f"Reviewed date updated for {m.get('doc_code')}.")
                        st.rerun()
                    except Exception as e:
                        st.error(f"Error saving: {e}")

                if is_obsolete:
                    if row[7].button("♻️", key=f"reactivate_{mid}", help="Reactivate this document"):
                        try:
                            sb.table("master_documents").update({"status": "Active"}).eq("id", mid).execute()
                            st.success(f"{code_txt} reactivated.")
                            st.rerun()
                        except Exception as e:
                            st.error(f"Error reactivating: {e}")
                else:
                    if row[7].button("🚫", key=f"obsolete_{mid}", help="Mark as obsolete (kept in the list, excluded from reviews)"):
                        try:
                            sb.table("master_documents").update({"status": "Obsolete"}).eq("id", mid).execute()
                            st.success(f"{code_txt} marked obsolete.")
                            st.rerun()
                        except Exception as e:
                            st.error(f"Error: {e}")

                if row[8].button("🗑️", key=f"del_{mid}", help="Permanently delete from master list"):
                    st.session_state[f"confirm_del_{mid}"] = True

                # Compact detail line: revision, approver, status, and a file link if uploaded
                detail_bits = []
                if m.get("revision_label"):
                    detail_bits.append(f"Rev {m['revision_label']}")
                if m.get("revision_date"):
                    detail_bits.append(f"revised {m['revision_date']}")
                if m.get("issue_date"):
                    detail_bits.append(f"issued {m['issue_date']}")
                if m.get("approved_by"):
                    detail_bits.append(f"approved by {m['approved_by']}")
                if m.get("status") and m.get("status") != "Active":
                    detail_bits.append(f"status: {m['status']}")
                detail_line = " • ".join(detail_bits)
                if detail_line or m.get("file_url"):
                    dcol1, dcol2 = st.columns([5,1])
                    if detail_line:
                        dcol1.caption(detail_line)
                    if m.get("file_url"):
                        dcol2.markdown(f"[📄 View file]({m['file_url']})")

                if st.session_state.get(f"confirm_del_{mid}"):
                    st.warning(
                        f"Delete **{m.get('doc_code')} — {m.get('title')}** from the master list? "
                        "This cannot be undone. Any document that references this code will keep "
                        "the reference, but it will no longer resolve to a master entry."
                    )
                    cc1, cc2 = st.columns([1,1])
                    with cc1:
                        if st.button("Yes, delete", key=f"confirm_yes_{mid}"):
                            try:
                                sb.table("master_documents").delete().eq("id", mid).execute()
                                st.session_state.pop(f"confirm_del_{mid}", None)
                                st.success("Deleted.")
                                st.rerun()
                            except Exception as e:
                                st.error(f"Error deleting: {e}")
                    with cc2:
                        if st.button("Cancel", key=f"confirm_no_{mid}"):
                            st.session_state.pop(f"confirm_del_{mid}", None)
                            st.rerun()
                st.markdown("<hr style='margin:4px 0;opacity:0.3'>", unsafe_allow_html=True)


# ─────────────────────────────────────────────
# IN-APP DOCUMENT RENDERER
# ─────────────────────────────────────────────
def _render_document(sb, doc, uid):
    doc_code  = doc.get("doc_code","—")
    title     = doc.get("title","")
    rev       = doc.get("revision", 0)
    revisions = doc.get("_revisions") or []

    st.markdown(f"""
<div style="border:1px solid #ccc;font-size:11px;margin-bottom:12px;border-radius:4px">
<table width="100%" style="border-collapse:collapse">
<tr>
  <td style="border:1px solid #ccc;padding:4px 8px;width:12%;font-weight:bold;background:#f5f5f5">TITLE</td>
  <td style="border:1px solid #ccc;padding:4px 8px;text-align:center;font-weight:bold">{title.upper()}</td>
  <td style="border:1px solid #ccc;padding:4px 8px;width:15%;font-weight:bold;background:#f5f5f5">NBR. OF PAGES</td>
  <td style="border:1px solid #ccc;padding:4px 8px;width:8%;text-align:center">—</td>
</tr>
<tr>
  <td style="border:1px solid #ccc;padding:4px 8px;font-weight:bold;background:#f5f5f5">CONTROL NBR.</td>
  <td style="border:1px solid #ccc;padding:4px 8px;text-align:center">{doc_code}</td>
  <td style="border:1px solid #ccc;padding:4px 8px;font-weight:bold;background:#f5f5f5">REVISION DATE</td>
  <td style="border:1px solid #ccc;padding:4px 8px;text-align:center">{revisions[-1]["revised_date"] if revisions else "—"}</td>
</tr>
</table>
</div>""", unsafe_allow_html=True)

    try:
        st.image(LOGO_PATH, width=120)
    except Exception:
        pass

    dept_label = doc.get("dept_label") or doc.get("dept","")
    st.markdown(f"<h3 style='text-align:center'>{dept_label.upper()}</h3>", unsafe_allow_html=True)
    st.markdown(f"<h2 style='text-align:center'>{title.upper()}</h2>", unsafe_allow_html=True)

    approvals = doc.get("approvals") or []
    if approvals:
        cols = st.columns(len(approvals)+1)
        labels = ["Department","Function","Signature","Date"]
        with cols[0]:
            for l in labels:
                st.markdown(f"**{l}**")
        for i, ap in enumerate(approvals):
            with cols[i+1]:
                st.markdown(ap.get("department",""))
                st.markdown(f"*{_clean_approval_label(ap.get('function',''))}*")
                st.markdown("&nbsp;")
                st.markdown("&nbsp;")

    adoption = doc.get("date_of_adoption","—")
    st.markdown(f"<p style='text-align:center'><b>Date of Adoption:</b> {adoption}</p>",
                unsafe_allow_html=True)


    st.markdown("#### REVISION HISTORY")
    if revisions:
        st.dataframe(pd.DataFrame([{
            "Rev": f"{r['revision']:02d}", "Date": r.get("revised_date",""),
            "Status": r.get("status",""), "Description": r.get("description","")
        } for r in revisions]), hide_index=True, use_container_width=True)

    st.markdown("---")
    st.markdown("### 1.0 Introduction")
    if doc.get("purpose"):
        st.markdown("**1.1 Purpose**")
        st.markdown(f"1.1.1 &nbsp; {doc['purpose']}", unsafe_allow_html=True)
    for i, p in enumerate(doc.get("policy") or [], 1):
        if i == 1: st.markdown("**1.2 Policy**")
        st.markdown(f"1.2.{i} &nbsp; {p}", unsafe_allow_html=True)
    for i, s in enumerate(doc.get("scope") or [], 1):
        if i == 1: st.markdown("**1.3 Scope of Application**")
        st.markdown(f"1.3.{i} &nbsp; {s}", unsafe_allow_html=True)
    for i, r in enumerate(doc.get("responsibilities") or [], 1):
        if i == 1: st.markdown("**1.4 Authorities & Responsibilities**")
        st.markdown(f"1.4.{i} &nbsp; {r}", unsafe_allow_html=True)

    st.markdown("---")
    st.markdown("### 2.0 Abbreviations and Definitions")
    abbrevs = doc.get("abbreviations") or []
    if abbrevs:
        st.dataframe(pd.DataFrame([{"Term": a.get("term",""),
                                    "Definition": a.get("definition","")}
                                   for a in abbrevs]),
                     hide_index=True, use_container_width=True)

    st.markdown("---")
    st.markdown("### 3.0 Procedure")
    steps = doc.get("procedure_steps") or []
    step_num = 0
    for step in steps:
        stype = step.get("type", "step")
        if stype == "input_row":
            items = step.get("items", [])
            titles = "; ".join(it.get("title", "") for it in items)
            heading = step.get("heading", "Trigger Sources")
            intro   = step.get("intro", "This process may be triggered by any of the following")
            st.markdown(f"**{heading}**")
            st.markdown(f"{intro}: {titles}.")
            continue
        if stype == "branch_split":
            for br in step.get("branches", []):
                st.markdown(f"**{br.get('label','')}**")
                for s in br.get("steps", []):
                    st.markdown(f"- {s.get('title','')}")
            continue
        step_num += 1
        st.markdown(f"**{step_num}. {step.get('title','')}**")
        if step.get("text"):
            st.markdown(step["text"])
        if step.get("side_branch"):
            st.markdown(f"→ *In parallel:* {step['side_branch'].get('title','')}")


    st.markdown("---")
    st.markdown("### 4.0 Associated Documentation")
    for refs, label in [
        (doc.get("related_docs") or [],      "4.1 Related Documents"),
        (doc.get("resulting_records") or [],  "4.2 Resulting Records"),
        (doc.get("ext_references") or [],     "4.3 Internal / External References"),
    ]:
        if refs:
            st.markdown(f"**{label}**")
            st.dataframe(pd.DataFrame([{"Code": r.get("code","—"),
                                        "Title": r.get("title","—")}
                                       for r in refs]),
                         hide_index=True, use_container_width=True)

    st.markdown(
        "<p style='font-size:10px;color:#aaa;text-align:center;margin-top:16px'>"
        "PROPRIETARY TO NAPCO NATIONAL — <b>UNCONTROLLED IF PRINTED</b></p>",
        unsafe_allow_html=True)

    if can_write():
        st.markdown("---")
        st.markdown("### Document Maintenance")

        # ── Reviewed-date status + Mark as Reviewed (no content change) ──
        reviewed_date = doc.get("reviewed_date")
        status_col, btn_col = st.columns([3, 1])
        with status_col:
            if reviewed_date:
                try:
                    rd = date.fromisoformat(str(reviewed_date)[:10])
                    next_review = rd.replace(year=rd.year + 1)
                    today = date.today()
                    if next_review <= today:
                        st.caption(f"🔴 Last reviewed {reviewed_date} — overdue (was due {next_review.isoformat()})")
                    elif (next_review - today).days <= 30:
                        st.caption(f"🟡 Last reviewed {reviewed_date} — next review due {next_review.isoformat()}")
                    else:
                        st.caption(f"🟢 Last reviewed {reviewed_date} — next review due {next_review.isoformat()}")
                except Exception:
                    st.caption(f"Last reviewed {reviewed_date}")
            else:
                st.caption("⚪ Never reviewed")
        with btn_col:
            if st.button("✅ Mark as Reviewed", key=f"reviewed_{doc['id']}",
                         help="Confirms this document is still accurate as-is — no content changes. "
                              "Updates the review date shown to auditors without creating a new revision."):
                try:
                    today_iso = date.today().isoformat()
                    sb.table("proc_documents").update({"reviewed_date": today_iso}).eq("id", doc["id"]).execute()
                    sb.table("master_documents").update({"reviewed_date": today_iso}).eq("doc_code", doc_code).execute()
                    st.success(f"Marked as reviewed on {today_iso}.")
                    st.rerun()
                except Exception as e:
                    st.error(f"Error: {e}")

        # ── Edit Document (actual content changes, logged as a new revision) ──
        with st.expander("✏️ Edit Document"):
            st.caption(
                "Changing anything here and saving creates a new revision, logs what changed, "
                "and marks the document as reviewed today. Procedure Steps are edited as JSON — "
                "ask Claude to help generate the updated structure if you're restructuring the "
                "flowchart (branches, side-branches, etc.), then paste it in below."
            )
            with st.form(f"edit_{doc['id']}"):
                e_title      = st.text_input("Title", value=doc.get("title", ""))
                e_dept_label = st.text_input("Department Label", value=doc.get("dept_label", ""))
                e_purpose    = st.text_area("Purpose", value=doc.get("purpose", ""), height=100)
                e_policy     = st.text_area("Policy (one per line)",
                                             value="\n".join(doc.get("policy") or []), height=80)
                e_scope      = st.text_area("Scope (one per line)",
                                             value="\n".join(doc.get("scope") or []), height=80)
                e_resp       = st.text_area("Responsibilities (one per line)",
                                             value="\n".join(doc.get("responsibilities") or []), height=120)

                abbr_lines = "\n".join(f"{a.get('term','')} = {a.get('definition','')}"
                                        for a in (doc.get("abbreviations") or []))
                e_abbr = st.text_area("Abbreviations (Term = Definition, one per line)",
                                       value=abbr_lines, height=100)

                steps_json = json.dumps(doc.get("procedure_steps") or [], indent=2, ensure_ascii=False)
                e_steps = st.text_area("Procedure Steps (JSON)", value=steps_json, height=300)

                related_lines = "\n".join(f"{r.get('code','')} | {r.get('title','')}"
                                           for r in (doc.get("related_docs") or []))
                e_related = st.text_area("Related Documents (code | title, one per line)",
                                          value=related_lines, height=80)

                resulting_lines = "\n".join(f"{r.get('code','')} | {r.get('title','')}"
                                             for r in (doc.get("resulting_records") or []))
                e_resulting = st.text_area("Resulting Records (code | title, one per line)",
                                            value=resulting_lines, height=80)

                ext_lines = "\n".join(f"{r.get('code','')} | {r.get('title','')}"
                                       for r in (doc.get("ext_references") or []))
                e_ext = st.text_area("External References (code | title, one per line)",
                                      value=ext_lines, height=80)

                approvals_lines = "\n".join(f"{a.get('department','')} | {a.get('function','')}"
                                             for a in (doc.get("approvals") or []))
                e_approvals = st.text_area("Approvals (department | function, one per line)",
                                            value=approvals_lines, height=80)

                e_rev_desc = st.text_area("What changed in this edit? *", height=80,
                                           help="Required — logged in the revision history.")

                submitted = st.form_submit_button("Save Changes")

            if submitted:
                if not e_rev_desc.strip():
                    st.error("Please describe what changed before saving.")
                else:
                    try:
                        steps_parsed = json.loads(e_steps)
                    except Exception as ex:
                        st.error(f"Procedure Steps JSON is invalid, nothing was saved: {ex}")
                        steps_parsed = None

                    if steps_parsed is not None:
                        def _parse_lines(text):
                            return [l.strip() for l in text.split("\n") if l.strip()]

                        def _parse_pairs(text):
                            out = []
                            for l in text.split("\n"):
                                l = l.strip()
                                if l and "|" in l:
                                    code, _, ttl = l.partition("|")
                                    out.append({"code": code.strip(), "title": ttl.strip()})
                            return out

                        def _parse_abbrevs(text):
                            out = []
                            for l in text.split("\n"):
                                l = l.strip()
                                if l and "=" in l:
                                    term, _, defn = l.partition("=")
                                    out.append({"term": term.strip(), "definition": defn.strip()})
                            return out

                        def _parse_approvals(text):
                            out = []
                            for l in text.split("\n"):
                                l = l.strip()
                                if l and "|" in l:
                                    dept, _, func = l.partition("|")
                                    out.append({"department": dept.strip(), "function": func.strip()})
                            return out

                        new_rev   = rev + 1
                        today_iso = date.today().isoformat()
                        payload = {
                            "title":             e_title,
                            "dept_label":        e_dept_label,
                            "purpose":           e_purpose,
                            "policy":            _parse_lines(e_policy),
                            "scope":             _parse_lines(e_scope),
                            "responsibilities":  _parse_lines(e_resp),
                            "abbreviations":     _parse_abbrevs(e_abbr),
                            "procedure_steps":   steps_parsed,
                            "related_docs":      _parse_pairs(e_related),
                            "resulting_records": _parse_pairs(e_resulting),
                            "ext_references":    _parse_pairs(e_ext),
                            "approvals":         _parse_approvals(e_approvals),
                            "revision":          new_rev,
                            "reviewed_date":     today_iso,
                            "updated_by":        uid,
                        }
                        try:
                            snapshot = {k: v for k, v in doc.items() if k != "_revisions"}
                            sb.table("proc_revisions").insert({
                                "doc_id":       doc["id"],
                                "revision":     new_rev,
                                "revised_date": today_iso,
                                "status":       revision_label(new_rev),
                                "description":  e_rev_desc,
                                "revised_by":   uid,
                                "snapshot":     snapshot,
                            }).execute()
                            sb.table("proc_documents").update(payload).eq("id", doc["id"]).execute()
                            sb.table("master_documents").update({
                                "title": e_title, "reviewed_date": today_iso,
                            }).eq("doc_code", doc_code).execute()
                            st.success(f"✅ Saved. Revision {new_rev:02d} created and reviewed date updated.")
                            st.rerun()
                        except Exception as e:
                            st.error(f"Error saving: {e}")


# ─────────────────────────────────────────────
# UI HELPERS
# ─────────────────────────────────────────────
def _edit_list(key, placeholder):
    items = st.session_state[key]
    for i, item in enumerate(items):
        c1, c2 = st.columns([8,1])
        with c1:
            st.session_state[key][i] = st.text_input(
                f"{placeholder} {i+1}", value=item, key=f"{key}_{i}",
                label_visibility="collapsed", placeholder=placeholder)
        with c2:
            if st.button("🗑️", key=f"{key}_del_{i}") and len(items) > 1:
                items.pop(i); st.rerun()
    if st.button(f"➕ Add {placeholder.lower()}", key=f"{key}_add"):
        st.session_state[key].append(""); st.rerun()


def _edit_refs(key, sb, uid):
    refs = st.session_state[key]
    for i, ref in enumerate(refs):
        c1, c2, c3 = st.columns([2,4,1])
        with c1:
            st.session_state[key][i]["code"] = st.text_input(
                "Code", value=ref.get("code",""), key=f"{key}_code_{i}",
                label_visibility="collapsed", placeholder="Doc code")
        with c2:
            st.session_state[key][i]["title"] = st.text_input(
                "Title", value=ref.get("title",""), key=f"{key}_title_{i}",
                label_visibility="collapsed", placeholder="Document title")
        with c3:
            if st.button("🗑️", key=f"{key}_del_{i}") and len(refs) > 1:
                refs.pop(i); st.rerun()
        code = ref.get("code","").strip()
        if code:
            exists = sb.table("master_documents").select("id").eq("doc_code", code).execute()
            if not exists.data:
                st.warning(f"⚠️ `{code}` not in master list — will be added on save.")
            else:
                st.success(f"✅ `{code}` found in master list.")
    if st.button("➕ Add reference", key=f"{key}_add"):
        st.session_state[key].append({"code":"","title":""}); st.rerun()


def _clear_form():
    for key in ["approvals","policy_points","scope_points","resp_points",
                "abbrevs","proc_steps","rel_docs","rec_docs","ext_refs"]:
        if key in st.session_state:
            del st.session_state[key]